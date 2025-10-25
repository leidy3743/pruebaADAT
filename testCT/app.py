import json
import os
import openai
import traceback
from flask import Flask, render_template, redirect, url_for, flash, request, jsonify, session, send_file
from flask_login import LoginManager, login_user, login_required, logout_user, UserMixin, current_user

from werkzeug.security import generate_password_hash, check_password_hash
# Utilidades y librerías estándar/faltantes
import io
from datetime import datetime, timedelta
import zipfile
from sqlalchemy import text, inspect
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import landscape, A4
import markdown
from reportlab.lib import colors

# Inicialización de la app Flask y LoginManager

app = Flask(__name__)
if os.path.exists('config.py'):
    app.config.from_object('config.Config')
else:
    app.config['SECRET_KEY'] = 'dev-secret-key'

# Registrar la extensión de SQLAlchemy
from models import db
db.init_app(app)

login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Importar modelos y utilidades desde models.py
from models import db, User, QuestionCuatro, ResultadoQuizCuatro, Curso, Nivel, Grado, Answer, QuestionTres, usuarios_cursos, nivel_por_grados, grados_dictados, Question, QuestionDos, ResultadoQuiz, ResultadoQuizDos, ResultadoQuizTres

# Inicializar Flask-Migrate
from flask_migrate import Migrate
migrate = Migrate(app, db)
from cache_utils import get_cached_colegios, get_cached_cursos, get_cached_niveles, get_cached_grados
# Importar formularios desde forms.py
from forms import RegistrationForm, LoginForm, QuizForm, QuizFormDos, QuizFormTres, QuizFormCuatro
# Importar IntegrityError para manejo de errores de integridad
from sqlalchemy.exc import IntegrityError


# Ruta para el test Marco Roman
@app.route('/quiz4', methods=['GET', 'POST'])
@login_required
def quiz4():
    # Flag de acceso al test PC
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='quiz4_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El Test de Pensamiento Computacional está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    resultado = ResultadoQuizCuatro.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "success")
        return redirect(url_for('quiz4_results', result_id=resultado.id))

    questions = QuestionCuatro.query.order_by(QuestionCuatro.id).all()

    if request.method == 'POST':
        user_answers = request.form.to_dict()
        correct_count = 0
        incorrect_count = 0
        score = 0
        for question in questions:
            user_answer = user_answers.get(str(question.id))
            if user_answer == question.correct_answer:
                correct_count += 1
                score += question.percentage if question.percentage else 1
            else:
                incorrect_count += 1
        resultado = ResultadoQuizCuatro(
            user_id=current_user.id,
            score=score,
            respuestas_correctas=correct_count,
            respuestas_incorrectas=incorrect_count
        )
        db.session.add(resultado)
        db.session.commit()
        return redirect(url_for('quiz4_results', result_id=resultado.id))

    return render_template('quiz4.html', questions=questions)

# Ruta para mostrar resultados de quiz4
@app.route('/quiz4_results/<int:result_id>')
def quiz4_results(result_id):
    resultado = ResultadoQuizCuatro.query.get_or_404(result_id)
    return render_template('quiz4_results.html', resultado=resultado)





@app.route('/', methods=['GET'])
def home():
    return render_template('home.html')  # Redirige a la ruta de registro


@app.route('/dashboard')
@login_required
def dashboard():
    print(f"=== DEBUG DASHBOARD ===")
    print(f"Usuario autenticado: {current_user.is_authenticated}")
    print(f"Username: {current_user.username if current_user.is_authenticated else 'N/A'}")
    print(f"Rol: {current_user.rol if current_user.is_authenticated else 'N/A'}")
    print(f"======================")
    return render_template('dashboard.html')


@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    # Flag de acceso a Perfil
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='perfil_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('La vista de Perfil está desactivada por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        try:
            # Actualizar nombre y cédula
            current_user.nombres = request.form.get('nombres')
            cedula = request.form.get('cedula')
            if cedula:
                current_user.cedula = cedula
            
            # Cambiar contraseña si se proporcionó
            new_password = request.form.get('new_password')
            confirm_password = request.form.get('confirm_password')
            
            if new_password:
                if new_password == confirm_password:
                    current_user.set_password(new_password)
                else:
                    flash('Las contraseñas no coinciden', 'error')
                    return render_template('profile.html')
            
            db.session.commit()
            flash('Perfil actualizado exitosamente', 'success')
            return redirect(url_for('profile'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar el perfil: {str(e)}', 'error')
            return render_template('profile.html')
    
    return render_template('profile.html')


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.route('/register', methods=['GET', 'POST'])
def register():
    print("Método de la solicitud:", request.method)

    import os
    print('=== DEBUG /register ===')
    print('SQLALCHEMY_DATABASE_URI:', app.config.get('SQLALCHEMY_DATABASE_URI'))
    colegios = get_cached_colegios()
    print('Colegios encontrados:', [(c.id, c.nombre) for c in colegios])
    form = RegistrationForm()
    form.colegio.choices = [(c.id, c.nombre) for c in colegios]
    form.cursos.query_factory = lambda: get_cached_cursos()
    form.nivel_grados.query_factory = lambda: get_cached_niveles()
    form.grados.query_factory = lambda: get_cached_grados()


    if request.method == "POST":
        print("Datos recibidos:", request.form)
        
        if form.validate():
            try:
                print("El formulario fue enviado y validado correctamente")
                
                # Extraer los IDs de las relaciones many-to-many
                cursos_ids = [g.id if isinstance(g, Curso) else g for g in form.cursos.data]
                grados_ids = [g.id if isinstance(g, Grado) else g for g in form.grados.data]
                nivel_ids = [g.id if isinstance(g, Nivel) else g for g in form.nivel_grados.data]
                
                # Crear el usuario SIN las relaciones many-to-many
                nuevo_usuario = User(
                    nombres=form.nombres.data,
                    correo=form.correo.data,
                    edad=form.edad.data,
                    cedula=form.cedula.data if form.cedula.data else None,
                    colegio_id=int(form.colegio.data),
                    institucion=form.institucion.data,
                    anios_experiencia=form.anios_experiencia.data,
                    nivel_educativo=form.nivel_educativo.data,
                    username=form.username.data,
                    password=generate_password_hash(request.form['password'], method='pbkdf2:sha256'),
                    rol=form.rol.data
                )

                # Ahora agregar las relaciones many-to-many
                if nivel_ids:
                    nuevo_usuario.nivel_grados = [db.session.get(Nivel, nid) for nid in nivel_ids if nid]
                
                if grados_ids:
                    nuevo_usuario.grados = [db.session.get(Grado, gid) for gid in grados_ids if gid]
                
                if cursos_ids:
                    nuevo_usuario.cursos = [db.session.get(Curso, cid) for cid in cursos_ids if cid]

                db.session.add(nuevo_usuario)
                db.session.commit()

                flash('¡Registro exitoso! Ahora puedes iniciar sesión.', 'success')
                return redirect(url_for('login'))

            except IntegrityError as e:
                db.session.rollback()  # Revertir cambios en caso de error

                if "Key (correo)" in str(e.orig):  # Detecta el error de correo duplicado
                    flash('El correo ya está registrado. Usa otro correo.', 'danger')
                else:
                    flash('Error al registrar usuario. Intenta nuevamente.', 'danger')

                print("Error de integridad en la base de datos:", str(e))
                traceback.print_exc()

            except Exception as e:
                db.session.rollback()
                print("Error al registrar usuario:", str(e))
                traceback.print_exc()
                flash(f'Error al registrar usuario: {str(e)}', 'danger')

        else:
            print("Errores en el formulario:", form.errors)
            for field, errors in form.errors.items():
                for error in errors:
                    flash(f'Error en {field}: {error}', 'danger')

    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    # Si el usuario ya está autenticado, redirigir al dashboard
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    print(f"=== LOGIN ROUTE CALLED - Method: {request.method} ===")
    form = LoginForm()
    if request.method == 'POST':
        print("POST REQUEST RECIBIDO")
        username = request.form['username']
        password = request.form['password']
        print(f"Usuario: {username}, Password: {password}")
        # Buscar usuario en la base de datos por email
        user = User.query.filter_by(username=username).first()
        
        if user:
            print(f"Usuario encontrado: {user.username}")
            print(f"Hash almacenado: {user.password}")
            print(f"Verificando contraseña...")
            password_match = check_password_hash(user.password, password)
            print(f"¿Contraseña coincide? {password_match}")
            
            if password_match:
                # Si el usuario existe y la contraseña es correcta
                print("user autenticado")
                login_user(user)
                # Actualizar fecha de último login
                user.last_login = datetime.utcnow()
                db.session.commit()
                # Limpiar mensajes flash anteriores y agregar solo el de bienvenida
                session.pop('_flashes', None)
                flash(f'¡Bienvenido/a {user.nombres}!', 'success')
                
                # Verificar si el usuario no tiene cédula registrada
                if not user.cedula:
                    session['show_cedula_modal'] = True
                
                return redirect(url_for('dashboard'))  # Redirige a la página principal después de iniciar sesión
            else:
                print("Contraseña incorrecta")
                flash('Usuario o contraseña incorrectos', 'danger')
        else:
            # Si el usuario no existe o la contraseña es incorrecta
            print(f"Usuario '{username}' no encontrado en la base de datos")
            flash('Usuario o contraseña incorrectos', 'danger')
    return render_template('login.html', form=form)


@app.route('/actualizar_cedula', methods=['POST'])
@login_required
def actualizar_cedula():
    """Actualizar cédula del usuario desde el modal"""
    cedula = request.form.get('cedula')
    if cedula:
        try:
            current_user.cedula = cedula
            db.session.commit()
            # Limpiar la flag del modal
            session.pop('show_cedula_modal', None)
            flash('Cédula actualizada exitosamente', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar cédula: {str(e)}', 'danger')
    return redirect(url_for('dashboard'))


@app.route('/logout')
@login_required
def logout():
    nombre_usuario = current_user.nombres
    logout_user()
    flash(f'Hasta pronto, {nombre_usuario}. Sesión cerrada correctamente.', 'info')
    return redirect(url_for('login'))


@app.route('/admin/users')
@login_required
def admin_users():
    """Ruta para ver todos los usuarios registrados"""
    usuarios = User.query.all()
    return render_template('admin_users.html', usuarios=usuarios)


@app.route('/gestion_usuarios')
@login_required
def gestion_usuarios():
    """Panel de gestión de usuarios (solo para administradores)"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    # Paginación para mejorar rendimiento
    page = request.args.get('page', 1, type=int)
    per_page = 20  # Mostrar 20 usuarios por página

    # Eager loading para evitar N+1 queries
    from sqlalchemy.orm import joinedload
    pagination = User.query.options(
        joinedload(User.colegio)
    ).order_by(User.id.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    # Obtener todos los roles únicos de la base de datos
    all_roles = [r[0] for r in db.session.query(User.rol).distinct().all() if r[0]]

    # DEBUG: Mostrar los primeros usuarios y su institucion
    print("======================")

    return render_template('gestion_usuarios.html', 
                         usuarios=pagination.items,
                         pagination=pagination,
                         all_roles=all_roles)


@app.route('/gestion_usuarios/crear', methods=['GET', 'POST'])
@login_required
def crear_usuario():
    """Crear nuevo usuario (solo admin)"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        try:
            nuevo_usuario = User(
                nombres=request.form.get('nombres'),
                correo=request.form.get('correo'),
                edad=int(request.form.get('edad')),
                username=request.form.get('username'),
                password=generate_password_hash(request.form.get('password'), method='pbkdf2:sha256'),
                nivel_educativo=request.form.get('nivel_educativo'),
                anios_experiencia=int(request.form.get('anios_experiencia')),
                rol=request.form.get('rol'),
                institucion=request.form.get('institucion')
            )
            db.session.add(nuevo_usuario)
            db.session.commit()
            flash('Usuario creado exitosamente', 'success')
            return redirect(url_for('gestion_usuarios'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al crear usuario: {str(e)}', 'danger')
    
    return render_template('crear_usuario.html')


@app.route('/gestion_usuarios/editar/<int:user_id>', methods=['GET', 'POST'])
@login_required
def editar_usuario(user_id):
    """Editar usuario existente (solo admin)"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    usuario = User.query.get_or_404(user_id)
    
    if request.method == 'POST':
        try:
            usuario.nombres = request.form.get('nombres')
            usuario.correo = request.form.get('correo')
            usuario.edad = int(request.form.get('edad'))
            usuario.cedula = request.form.get('cedula') if request.form.get('cedula') else None
            usuario.username = request.form.get('username')
            usuario.nivel_educativo = request.form.get('nivel_educativo')
            usuario.anios_experiencia = int(request.form.get('anios_experiencia'))
            usuario.rol = request.form.get('rol')
            usuario.institucion = request.form.get('institucion')
            
            # Actualizar colegio si se proporciona
            colegio_id = request.form.get('colegio_id')
            if colegio_id:
                usuario.colegio_id = int(colegio_id) if colegio_id != '' else None
            
            # Cambiar contraseña solo si se proporciona
            new_password = request.form.get('password')
            if new_password:
                usuario.password = generate_password_hash(new_password, method='pbkdf2:sha256')
            
            db.session.commit()
            flash('Usuario actualizado exitosamente', 'success')
            return redirect(url_for('gestion_usuarios'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar usuario: {str(e)}', 'danger')
    
    # Obtener colegios desde caché (ordenados alfabéticamente)
    colegios = get_cached_colegios()
    return render_template('editar_usuario.html', usuario=usuario, colegios=colegios)


@app.route('/gestion_usuarios/eliminar/<int:user_id>', methods=['POST'])
@login_required
def eliminar_usuario(user_id):
    """Eliminar usuario (solo admin)"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        usuario = User.query.get_or_404(user_id)
        
        # No permitir eliminar el propio usuario admin
        if usuario.id == current_user.id:
            flash('No puedes eliminar tu propio usuario', 'danger')
            return redirect(url_for('gestion_usuarios'))

        # 1) Eliminar dependencias explícitas para evitar violaciones de FK (NotNullViolation)
        # Resultados de quizzes
        try:
            ResultadoQuiz.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass
        try:
            ResultadoQuizDos.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass
        try:
            ResultadoQuizTres.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass
        try:
            ResultadoQuizCuatro.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass

        # Respuestas de quizzes
        try:
            Answer.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass

        # Actividades generadas
        try:
            ActividadGenerada.query.filter_by(user_id=usuario.id).delete(synchronize_session=False)
        except Exception:
            pass

        # 2) Limpiar tablas de asociación many-to-many
        try:
            db.session.execute(usuarios_cursos.delete().where(usuarios_cursos.c.user_id == usuario.id))
        except Exception:
            pass
        try:
            db.session.execute(nivel_por_grados.delete().where(nivel_por_grados.c.user_id == usuario.id))
        except Exception:
            pass
        try:
            db.session.execute(grados_dictados.delete().where(grados_dictados.c.user_id == usuario.id))
        except Exception:
            pass

        # 3) Finalmente eliminar el usuario
        db.session.delete(usuario)
        db.session.commit()
        flash('Usuario eliminado exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar usuario: {str(e)}', 'danger')
    
    return redirect(url_for('gestion_usuarios'))


# ========== REINICIO DE TESTS POR USUARIO (ADMIN) ==========

@app.route('/gestion_usuarios/reset/<int:user_id>/<quiz>', methods=['POST'])
@login_required
def reset_test_usuario(user_id, quiz):
    """Permite a un admin reiniciar un test específico o todos para un usuario dado.
    quiz puede ser: 'quiz1', 'quiz2', 'quiz3', 'quiz4' o 'all'.
    """
    if current_user.rol != 'admin':
        flash('No tienes permisos para realizar esta acción', 'danger')
        return redirect(url_for('dashboard'))

    usuario = User.query.get_or_404(user_id)

    if quiz.lower() == 'all':
        print(f"[DEPURACION] Reiniciando TODOS los tests para el usuario: {usuario.username} (ID: {usuario.id})")

    # Banderas para feedback
    acciones = []

    try:
        objetivo = quiz.lower()

        if objetivo in ('quiz1', 'all'):
            r1 = ResultadoQuiz.query.filter_by(user_id=usuario.id).first()
            if r1:
                db.session.delete(r1)
                acciones.append('Quiz 1 (ADAT)')

        if objetivo in ('quiz2', 'all'):
            r2 = ResultadoQuizDos.query.filter_by(user_id=usuario.id).first()
            if r2:
                db.session.delete(r2)
                acciones.append('Quiz 2 (Estilos)')

        if objetivo in ('quiz3', 'all'):
            r3 = ResultadoQuizTres.query.filter_by(user_id=usuario.id).first()
            if r3:
                db.session.delete(r3)
                acciones.append('Quiz 3 (Tipo de Jugador)')

        if objetivo in ('quiz4', 'all'):
            r4 = ResultadoQuizCuatro.query.filter_by(user_id=usuario.id).first()
            if r4:
                db.session.delete(r4)
                acciones.append('Quiz 4 (Pensamiento Computacional)')

        db.session.commit()

        # Verificar después de commit
        if objetivo in ('quiz4', 'all'):
            r4_post = ResultadoQuizCuatro.query.filter_by(user_id=usuario.id).first()

        if acciones:
            if objetivo == 'all':
                flash(f"Se reiniciaron todos los tests para el usuario {usuario.username}.", 'success')
            else:
                flash(f"Se reinició {', '.join(acciones)} para el usuario {usuario.username}.", 'success')
        else:
            if objetivo == 'all':
                flash(f"El usuario {usuario.username} no tenía resultados para reiniciar.", 'info')
            else:
                flash(f"El usuario {usuario.username} no tenía resultados del {objetivo}.", 'info')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al reiniciar tests para {usuario.username}: {str(e)}', 'danger')

    return redirect(url_for('gestion_usuarios'))


# ==================== GENERADOR DE CERTIFICADOS ====================

@app.route('/admin/generar_certificado/<int:user_id>')
@login_required
def generar_certificado(user_id):
    """Generar certificado PDF para un usuario específico"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    usuario = User.query.get_or_404(user_id)
    
    # Crear buffer para el PDF
    buffer = io.BytesIO()
    
    # Crear PDF en orientación horizontal (landscape)
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    
    # Colores elegantes
    color_principal = colors.HexColor('#1a237e')  # Azul oscuro
    color_acento = colors.HexColor('#0d47a1')     # Azul medio
    color_dorado = colors.HexColor('#b8860b')     # Dorado
    
    # ===== BORDE DECORATIVO =====
    c.setStrokeColor(color_dorado)
    c.setLineWidth(3)
    c.rect(30, 30, width-60, height-60, stroke=1, fill=0)
    
    c.setStrokeColor(color_principal)
    c.setLineWidth(1)
    c.rect(40, 40, width-80, height-80, stroke=1, fill=0)
    
    # ===== ENCABEZADO =====
    c.setFont("Helvetica-Bold", 28)
    c.setFillColor(color_principal)
    c.drawCentredString(width/2, height-100, "CERTIFICADO DE PARTICIPACIÓN")
    
    # Línea decorativa debajo del título
    c.setStrokeColor(color_dorado)
    c.setLineWidth(2)
    c.line(width/2-200, height-115, width/2+200, height-115)
    
    # ===== TEXTO PRINCIPAL =====
    c.setFont("Helvetica", 14)
    c.setFillColor(colors.black)
    c.drawCentredString(width/2, height-160, "Se otorga el presente certificado a:")
    
    # Nombre del usuario (destacado)
    c.setFont("Helvetica-Bold", 32)
    c.setFillColor(color_acento)
    c.drawCentredString(width/2, height-220, usuario.nombres.upper())
    
    # Línea decorativa debajo del nombre
    c.setStrokeColor(color_dorado)
    c.setLineWidth(1)
    c.line(width/2-250, height-235, width/2+250, height-235)
    
    # ===== CUERPO DEL CERTIFICADO =====
    c.setFont("Helvetica", 13)
    c.setFillColor(colors.black)
    
    # Texto de participación (justificado visualmente)
    texto_linea1 = "Por su destacada participación en la plataforma ADAT"
    texto_linea2 = "(Análisis de Datos para el Aprendizaje y Tecnología),"
    texto_linea3 = "completando exitosamente las evaluaciones de diagnóstico y"
    texto_linea4 = "contribuyendo al fortalecimiento de competencias educativas digitales."
    
    y_pos = height - 290
    c.drawCentredString(width/2, y_pos, texto_linea1)
    c.drawCentredString(width/2, y_pos-25, texto_linea2)
    c.drawCentredString(width/2, y_pos-50, texto_linea3)
    c.drawCentredString(width/2, y_pos-75, texto_linea4)
    
    # ===== INFORMACIÓN ADICIONAL =====
    c.setFont("Helvetica", 11)
    c.setFillColor(colors.grey)
    
    # Institución
    if usuario.institucion:
        c.drawCentredString(width/2, height-420, f"Institución: {usuario.institucion}")
    
    # Nivel educativo
    if usuario.nivel_educativo:
        nivel_texto = usuario.nivel_educativo.capitalize()
        c.drawCentredString(width/2, height-440, f"Nivel Educativo: {nivel_texto}")
    
    # ===== FECHA =====
    fecha_actual = datetime.now().strftime("%d de %B de %Y")
    # Traducir mes al español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    for eng, esp in meses.items():
        fecha_actual = fecha_actual.replace(eng, esp)
    
    c.setFont("Helvetica-Oblique", 12)
    c.setFillColor(colors.black)
    c.drawCentredString(width/2, 120, f"Fecha de emisión: {fecha_actual}")
    
    # ===== FIRMA/SELLO (decorativo) =====
    c.setFont("Helvetica", 10)
    c.setFillColor(colors.grey)
    
    # Línea de firma izquierda
    c.line(150, 160, 300, 160)
    c.drawCentredString(225, 145, "Administrador ADAT")
    
    # Línea de firma derecha
    c.line(width-300, 160, width-150, 160)
    c.drawCentredString(width-225, 145, "Plataforma ADAT")
    
    # ===== PIE DE PÁGINA =====
    c.setFont("Helvetica", 8)
    c.setFillColor(colors.lightgrey)
    c.drawCentredString(width/2, 70, "Este certificado ha sido generado digitalmente por la plataforma ADAT")
    c.drawCentredString(width/2, 55, f"ID de verificación: ADAT-{usuario.id:05d}-{datetime.now().strftime('%Y%m%d')}")
    
    # Guardar PDF
    c.save()
    
    # Preparar descarga
    buffer.seek(0)
    nombre_archivo = f"certificado_{usuario.username}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=nombre_archivo
    )


# ==================== GESTIÓN DE CONTENIDO DE TESTS ====================

@app.route('/gestion_tests')
@login_required
def gestion_tests():
    """Panel principal de gestión de tests"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    # Contar preguntas de cada test
    total_quiz1 = Question.query.count()
    total_quiz2 = QuestionDos.query.count()
    total_quiz3 = QuestionTres.query.count()
    total_quiz4 = QuestionCuatro.query.count()
    
    return render_template('gestion_tests.html', 
                         total_quiz1=total_quiz1,
                         total_quiz2=total_quiz2,
                         total_quiz3=total_quiz3,
                         total_quiz4=total_quiz4)


# ========== GESTIÓN QUIZ 1 (ADAT) ==========

@app.route('/gestion_tests/quiz1')
@login_required
def gestion_quiz1():
    """Ver todas las preguntas del Quiz 1 (ADAT)"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    preguntas = Question.query.all()
    return render_template('gestion_quiz1.html', preguntas=preguntas)


@app.route('/gestion_tests/quiz1/editar/<int:question_id>', methods=['GET', 'POST'])
@login_required
def editar_quiz1(question_id):
    """Editar pregunta del Quiz 1"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    pregunta = Question.query.get_or_404(question_id)
    form = QuizForm(obj=pregunta)
    
    if form.validate_on_submit():
        try:
            original_text = form.statement.data or ''
            pregunta.statement = original_text
            pregunta.option_a = form.option_a.data
            pregunta.option_b = form.option_b.data
            pregunta.option_c = form.option_c.data
            pregunta.option_d = form.option_d.data
            pregunta.correct_answer = form.correct_answer.data
            pregunta.label = form.label.data
            pregunta.percentage = form.percentage.data
            pregunta.image_url = form.image_url.data if form.image_url.data else ''

            # Intento 1: flush y verificar que no haya truncamiento por tipo de columna
            db.session.flush()
            saved_len = len(pregunta.statement or '')
            if saved_len < len(original_text):
                # Truncamiento detectado: intentar convertir columna a TEXT y reintentar
                try:
                    table_name = pregunta.__table__.name
                    dialect = db.engine.name
                    if dialect == 'postgresql':
                        db.session.execute(text(f'ALTER TABLE {table_name} ALTER COLUMN statement TYPE TEXT'))
                        db.session.flush()
                        # Reasignar y volver a intentar
                        pregunta.statement = original_text
                        db.session.flush()
                        saved_len = len(pregunta.statement or '')
                except Exception:
                    pass

            if saved_len < len(original_text):
                # Aún truncado: revertir y reportar
                db.session.rollback()
                flash('El enunciado es más largo que el tamaño permitido en la base de datos. He intentado ajustar la columna automáticamente sin éxito. Avísame y lo corrijo manualmente.', 'danger')
                return render_template('editar_quiz1.html', form=form, pregunta=pregunta)

            # Si todo bien, commit final
            db.session.commit()
            flash('Pregunta actualizada exitosamente', 'success')
            return redirect(url_for('gestion_quiz1'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar pregunta: {str(e)}', 'danger')
    # Si es POST pero la validación falló, mostrar errores para que el usuario sepa qué corregir
    if request.method == 'POST' and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f'{field}: {err}', 'danger')

    return render_template('editar_quiz1.html', form=form, pregunta=pregunta)


@app.route('/gestion_tests/quiz1/eliminar/<int:question_id>', methods=['POST'])
@login_required
def eliminar_quiz1(question_id):
    """Eliminar pregunta del Quiz 1"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        pregunta = Question.query.get_or_404(question_id)
        db.session.delete(pregunta)
        db.session.commit()
        flash('Pregunta eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar pregunta: {str(e)}', 'danger')
    
    return redirect(url_for('gestion_quiz1'))


# ========== GESTIÓN QUIZ 2 (Estilos de Aprendizaje) ==========

@app.route('/gestion_tests/quiz2')
@login_required
def gestion_quiz2():
    """Ver todas las preguntas del Quiz 2"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    preguntas = QuestionDos.query.all()
    return render_template('gestion_quiz2.html', preguntas=preguntas)


@app.route('/gestion_tests/quiz2/editar/<int:question_id>', methods=['GET', 'POST'])
@login_required
def editar_quiz2(question_id):
    """Editar pregunta del Quiz 2"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    pregunta = QuestionDos.query.get_or_404(question_id)
    form = QuizFormDos(obj=pregunta)
    
    if form.validate_on_submit():
        try:
            pregunta.statement = form.statement.data
            pregunta.option_a = form.option_a.data
            pregunta.option_b = form.option_b.data
            
            db.session.commit()
            flash('Pregunta actualizada exitosamente', 'success')
            return redirect(url_for('gestion_quiz2'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar pregunta: {str(e)}', 'danger')
    if request.method == 'POST' and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f'{field}: {err}', 'danger')

    return render_template('editar_quiz2.html', form=form, pregunta=pregunta)


@app.route('/gestion_tests/quiz2/eliminar/<int:question_id>', methods=['POST'])
@login_required
def eliminar_quiz2(question_id):
    """Eliminar pregunta del Quiz 2"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        pregunta = QuestionDos.query.get_or_404(question_id)
        db.session.delete(pregunta)
        db.session.commit()
        flash('Pregunta eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar pregunta: {str(e)}', 'danger')
    
    return redirect(url_for('gestion_quiz2'))


# ========== GESTIÓN QUIZ 3 (Tipo de Jugador) ==========

@app.route('/gestion_tests/quiz3')
@login_required
def gestion_quiz3():
    """Ver todas las preguntas del Quiz 3"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    preguntas = QuestionTres.query.all()
    return render_template('gestion_quiz3.html', preguntas=preguntas)


@app.route('/gestion_tests/quiz3/editar/<int:question_id>', methods=['GET', 'POST'])
@login_required
def editar_quiz3(question_id):
    """Editar pregunta del Quiz 3"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    pregunta = QuestionTres.query.get_or_404(question_id)
    form = QuizFormTres(obj=pregunta)
    
    if form.validate_on_submit():
        try:
            pregunta.statement = form.statement.data
            pregunta.option_a = form.option_a.data
            pregunta.option_b = form.option_b.data
            pregunta.option_c = form.option_c.data
            pregunta.option_d = form.option_d.data
            pregunta.option_e = form.option_e.data
            
            db.session.commit()
            flash('Pregunta actualizada exitosamente', 'success')
            return redirect(url_for('gestion_quiz3'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar pregunta: {str(e)}', 'danger')
    if request.method == 'POST' and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f'{field}: {err}', 'danger')

    return render_template('editar_quiz3.html', form=form, pregunta=pregunta)


@app.route('/gestion_tests/quiz3/eliminar/<int:question_id>', methods=['POST'])
@login_required
def eliminar_quiz3(question_id):
    """Eliminar pregunta del Quiz 3"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        pregunta = QuestionTres.query.get_or_404(question_id)
        db.session.delete(pregunta)
        db.session.commit()
        flash('Pregunta eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar pregunta: {str(e)}', 'danger')
    
    return redirect(url_for('gestion_quiz3'))


# ========== GESTIÓN QUIZ 4 (Pensamiento Computacional) ==========

@app.route('/gestion_tests/quiz4')
@login_required
def gestion_quiz4():
    """Ver todas las preguntas del Quiz 4"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    preguntas = QuestionCuatro.query.all()
    return render_template('gestion_quiz4.html', preguntas=preguntas)


@app.route('/gestion_tests/quiz4/editar/<int:question_id>', methods=['GET', 'POST'])
@login_required
def editar_quiz4(question_id):
    """Editar pregunta del Quiz 4"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    pregunta = QuestionCuatro.query.get_or_404(question_id)
    form = QuizFormCuatro(obj=pregunta)
    
    if form.validate_on_submit():
        try:
            pregunta.statement = form.statement.data
            pregunta.option_a = form.option_a.data
            pregunta.option_b = form.option_b.data
            pregunta.option_c = form.option_c.data
            pregunta.option_d = form.option_d.data
            pregunta.correct_answer = form.correct_answer.data
            pregunta.label = form.label.data
            pregunta.percentage = form.percentage.data
            # Preferir imagen subida; si no, usar URL normalizada
            uploaded = request.files.get('image_file')
            pregunta.image_url = form.image_url.data if form.image_url.data else ''

            db.session.commit()
            flash('Pregunta actualizada exitosamente', 'success')
            return redirect(url_for('gestion_quiz4'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar pregunta: {str(e)}', 'danger')
    if request.method == 'POST' and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f'{field}: {err}', 'danger')

    return render_template('editar_quiz4.html', form=form, pregunta=pregunta)


@app.route('/gestion_tests/quiz4/eliminar/<int:question_id>', methods=['POST'])
@login_required
def eliminar_quiz4(question_id):
    """Eliminar pregunta del Quiz 4"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        pregunta = QuestionCuatro.query.get_or_404(question_id)
        db.session.delete(pregunta)
        db.session.commit()
        flash('Pregunta eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar pregunta: {str(e)}', 'danger')
    
    return redirect(url_for('gestion_quiz4'))


# ==================== ESTADÍSTICAS PARA ADMIN ====================

@app.route('/admin/estadisticas')
@login_required
def estadisticas():
    """Dashboard de estadísticas completo para administradores"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    from sqlalchemy import func
    from datetime import datetime, timedelta
    from cache_utils import cache_get_or_set
    
    # Cachear todas las métricas 5 minutos para aliviar carga de agregaciones
    def _compute_stats():
        # === ESTADÍSTICAS GENERALES ===
        total_usuarios = User.query.count()

        # Completitud por test
        quiz1_completados = ResultadoQuiz.query.count()
        quiz2_completados = ResultadoQuizDos.query.count()
        quiz3_completados = ResultadoQuizTres.query.count()
        quiz4_completados = ResultadoQuizCuatro.query.count()

        total_tests_completados = quiz1_completados + quiz2_completados + quiz3_completados + quiz4_completados

        # Porcentajes de completitud
        quiz1_porcentaje = round((quiz1_completados / total_usuarios * 100) if total_usuarios > 0 else 0, 1)
        quiz2_porcentaje = round((quiz2_completados / total_usuarios * 100) if total_usuarios > 0 else 0, 1)
        quiz3_porcentaje = round((quiz3_completados / total_usuarios * 100) if total_usuarios > 0 else 0, 1)
        quiz4_porcentaje = round((quiz4_completados / total_usuarios * 100) if total_usuarios > 0 else 0, 1)

        # === ESTADÍSTICAS QUIZ 4 (Pensamiento Computacional) ===
        quiz4_promedio = db.session.query(func.avg(ResultadoQuizCuatro.score)).scalar() or 0
        quiz4_promedio = round(quiz4_promedio, 1)

        quiz4_max = db.session.query(func.max(ResultadoQuizCuatro.score)).scalar() or 0
        quiz4_min = db.session.query(func.min(ResultadoQuizCuatro.score)).scalar() or 0

        # Top 10 usuarios en Quiz 4 por respuestas correctas
        top_quiz4 = (
            db.session.query(
                User.nombres, 
                ResultadoQuizCuatro.score,
                ResultadoQuizCuatro.respuestas_correctas
            )
            .join(ResultadoQuizCuatro)
            .order_by(ResultadoQuizCuatro.respuestas_correctas.desc(), ResultadoQuizCuatro.score.desc())
            .limit(10)
            .all()
        )

        # === ESTADÍSTICAS POR INSTITUCIÓN ===
        stats_institucion = db.session.query(
            User.colegio_id,
            func.count(User.id).label('total_usuarios')
    ).group_by(User.colegio_id).order_by(func.count(User.id).desc()).all()

        # === ESTADÍSTICAS POR NIVEL EDUCATIVO ===
        stats_nivel = db.session.query(
            User.nivel_educativo,
            func.count(User.id).label('total')
        ).group_by(User.nivel_educativo).order_by(func.count(User.id).desc()).all()

        # === ESTADÍSTICAS POR ROL ===
        stats_rol = db.session.query(
            User.rol,
            func.count(User.id).label('total')
        ).group_by(User.rol).all()

        # === DISTRIBUCIÓN DE EDADES ===
        edad_promedio = db.session.query(func.avg(User.edad)).scalar() or 0
        edad_promedio = round(edad_promedio, 1)

        # Rangos de edad
        rango_18_25 = User.query.filter(User.edad >= 18, User.edad <= 25).count()
        rango_26_35 = User.query.filter(User.edad >= 26, User.edad <= 35).count()
        rango_36_45 = User.query.filter(User.edad >= 36, User.edad <= 45).count()
        rango_46_mas = User.query.filter(User.edad >= 46).count()

        # === USUARIOS RECIENTES ===
        usuarios_recientes = User.query.filter(User.last_login.isnot(None)).order_by(User.last_login.desc()).limit(10).all()

        # === RESULTADOS RECIENTES ===
        resultados_recientes_quiz4 = db.session.query(
            User.nombres,
            ResultadoQuizCuatro.score,
            ResultadoQuizCuatro.id
        ).join(ResultadoQuizCuatro).order_by(ResultadoQuizCuatro.id.desc()).limit(5).all()

        # === USUARIOS Y TESTS COMPLETADOS ===
        # Obtener parámetros de filtro
        buscar_nombre = request.args.get('buscar_nombre', '').strip()
        filtro_test1 = request.args.get('filtro_test1', '')
        filtro_test2 = request.args.get('filtro_test2', '')
        filtro_test3 = request.args.get('filtro_test3', '')
        filtro_test4 = request.args.get('filtro_test4', '')
        
        # Consulta base
        query = db.session.query(
            User.nombres,
            User.username,
            User.institucion,
            User.nivel_educativo,
            User.rol,
            func.coalesce(ResultadoQuiz.id, 0).label('quiz1_completado'),
            func.coalesce(ResultadoQuizDos.id, 0).label('quiz2_completado'),
            func.coalesce(ResultadoQuizTres.id, 0).label('quiz3_completado'),
            func.coalesce(ResultadoQuizCuatro.id, 0).label('quiz4_completado')
        ).outerjoin(ResultadoQuiz, User.id == ResultadoQuiz.user_id)\
         .outerjoin(ResultadoQuizDos, User.id == ResultadoQuizDos.user_id)\
         .outerjoin(ResultadoQuizTres, User.id == ResultadoQuizTres.user_id)\
         .outerjoin(ResultadoQuizCuatro, User.id == ResultadoQuizCuatro.user_id)
        
        # Aplicar filtros
        if buscar_nombre:
            query = query.filter(User.nombres.ilike(f'%{buscar_nombre}%'))
        
        if filtro_test1 == 'completado':
            query = query.filter(ResultadoQuiz.id.isnot(None))
        elif filtro_test1 == 'no_completado':
            query = query.filter(ResultadoQuiz.id.is_(None))
            
        if filtro_test2 == 'completado':
            query = query.filter(ResultadoQuizDos.id.isnot(None))
        elif filtro_test2 == 'no_completado':
            query = query.filter(ResultadoQuizDos.id.is_(None))
            
        if filtro_test3 == 'completado':
            query = query.filter(ResultadoQuizTres.id.isnot(None))
        elif filtro_test3 == 'no_completado':
            query = query.filter(ResultadoQuizTres.id.is_(None))
            
        if filtro_test4 == 'completado':
            query = query.filter(ResultadoQuizCuatro.id.isnot(None))
        elif filtro_test4 == 'no_completado':
            query = query.filter(ResultadoQuizCuatro.id.is_(None))
        
        usuarios_tests = query.order_by(User.nombres).all()

        return {
            'total_usuarios': total_usuarios,
            'quiz1_completados': quiz1_completados,
            'quiz2_completados': quiz2_completados,
            'quiz3_completados': quiz3_completados,
            'quiz4_completados': quiz4_completados,
            'total_tests_completados': total_tests_completados,
            'quiz1_porcentaje': quiz1_porcentaje,
            'quiz2_porcentaje': quiz2_porcentaje,
            'quiz3_porcentaje': quiz3_porcentaje,
            'quiz4_porcentaje': quiz4_porcentaje,
            'quiz4_promedio': quiz4_promedio,
            'quiz4_max': quiz4_max,
            'quiz4_min': quiz4_min,
            'top_quiz4': top_quiz4,
            'stats_institucion': stats_institucion,
            'stats_nivel': stats_nivel,
            'stats_rol': stats_rol,
            'edad_promedio': edad_promedio,
            'rango_18_25': rango_18_25,
            'rango_26_35': rango_26_35,
            'rango_36_45': rango_36_45,
            'rango_46_mas': rango_46_mas,
            'usuarios_recientes': usuarios_recientes,
            'resultados_recientes_quiz4': resultados_recientes_quiz4,
            'usuarios_tests': usuarios_tests,
        }

    stats = cache_get_or_set('admin_estadisticas', _compute_stats, ttl=300)
    
    return render_template('admin_estadisticas.html', **stats)


# ==================== EXPORTACIÓN DE DATOS ====================

@app.route('/admin/exportar/usuarios')
@login_required
def exportar_usuarios():
    """Exportar todos los usuarios a Excel"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import send_file
    
    # Crear workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Usuarios"
    
    # Encabezados
    headers = ['ID', 'Nombres', 'Usuario', 'Correo', 'Edad', 'Institución', 
               'Nivel Educativo', 'Años Experiencia', 'Rol']
    ws.append(headers)
    
    # Estilo de encabezados
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Obtener usuarios
    usuarios = User.query.all()
    
    # Agregar datos
    for usuario in usuarios:
        ws.append([
            usuario.id,
            usuario.nombres,
            usuario.username,
            usuario.correo,
            usuario.edad,
            usuario.institucion,
            usuario.nivel_educativo,
            usuario.anios_experiencia,
            usuario.rol
        ])
    
    # Ajustar ancho de columnas
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Guardar en memoria
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    from datetime import datetime
    fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'usuarios_adat_{fecha}.xlsx'
    )


@app.route('/admin/exportar/resultados/<quiz>')
@login_required
def exportar_resultados(quiz):
    """Exportar resultados de un quiz específico a Excel"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import send_file
    from datetime import datetime
    
    wb = Workbook()
    ws = wb.active
    
    # Estilo de encabezados
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    if quiz == 'quiz1':
        ws.title = "Resultados ADAT"
        headers = ['ID Usuario', 'Nombres', 'Usuario', 'Institución', 'Puntuaciones (JSON)']
        ws.append(headers)
        
        resultados = db.session.query(
            User.id, User.nombres, User.username, User.colegio_id, ResultadoQuiz.score
        ).join(ResultadoQuiz).all()
        
        for resultado in resultados:
            ws.append([resultado[0], resultado[1], resultado[2], resultado[3], str(resultado[4])])
    
    elif quiz == 'quiz2':
        ws.title = "Resultados Estilos Aprendizaje"
        headers = ['ID Usuario', 'Nombres', 'Usuario', 'Institución', 
                   'Sensorial-Intuitivo', 'Visual-Verbal', 'Activo-Reflexivo', 'Secuencial-Global']
        ws.append(headers)
        
        resultados = db.session.query(
            User.id, User.nombres, User.username, User.colegio_id,
            ResultadoQuizDos.sensorial_intuitivo,
            ResultadoQuizDos.visual_verbal,
            ResultadoQuizDos.activo_reflexivo,
            ResultadoQuizDos.secuencial_global
        ).join(ResultadoQuizDos).all()
        
        for resultado in resultados:
            ws.append([resultado[0], resultado[1], resultado[2], resultado[3],
                      str(resultado[4]), str(resultado[5]), str(resultado[6]), str(resultado[7])])
    
    elif quiz == 'quiz3':
        ws.title = "Resultados Tipo Jugador"
        headers = ['ID Usuario', 'Nombres', 'Usuario', 'Institución',
                   'Filántropo', 'Socializador', 'Triunfador', 'Jugador', 'Espíritu Libre', 'Disruptor']
        ws.append(headers)
        
        resultados = db.session.query(
            User.id, User.nombres, User.username, User.colegio_id,
            ResultadoQuizTres.filantropo, ResultadoQuizTres.socializador,
            ResultadoQuizTres.triunfador, ResultadoQuizTres.jugador,
            ResultadoQuizTres.espiritu_libre, ResultadoQuizTres.disruptor
        ).join(ResultadoQuizTres).all()
        
        for resultado in resultados:
            ws.append(list(resultado))
    
    elif quiz == 'quiz4':
        ws.title = "Resultados Pensamiento Comp"
        headers = ['ID Usuario', 'Nombres', 'Usuario', 'Institución', 
                   'Puntuación', 'Correctas', 'Incorrectas']
        ws.append(headers)
        
        resultados = db.session.query(
            User.id, User.nombres, User.username, User.colegio_id,
            ResultadoQuizCuatro.score, ResultadoQuizCuatro.respuestas_correctas,
            ResultadoQuizCuatro.incorrect_count
        ).join(ResultadoQuizCuatro).all()
        
        for resultado in resultados:
            ws.append(list(resultado))
    
    else:
        flash('Quiz no válido', 'danger')
        return redirect(url_for('estadisticas'))
    
    # Aplicar estilo a encabezados
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Ajustar ancho de columnas
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Guardar en memoria
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'resultados_{quiz}_{fecha}.xlsx'
    )


# ==================== BACKUP COMPLETO DE BASE DE DATOS ====================

@app.route('/admin/backup', methods=['GET'])
@login_required
def admin_backup():
    """Genera una copia completa de la base de datos.
    - Si es SQLite con archivo local: descarga el .db directamente.
    - Para otros motores: genera un backup JSON de todas las tablas (esquema + datos) comprimido en ZIP.
    """
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))

    from urllib.parse import urlparse
    from datetime import datetime as _dt
    from sqlalchemy import MetaData, Table, select

    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    ts = _dt.now().strftime('%Y%m%d_%H%M%S')

    # Caso 1: SQLite archivo local -> entregar el archivo .db
    try:
        if db_url.startswith('sqlite:///') or db_url.startswith('sqlite:////'):
            # Normalizar ruta de archivo
            path = None
            if db_url.startswith('sqlite:////'):
                path = '/' + db_url.replace('sqlite:////', '', 1)
            else:  # sqlite:///
                path = db_url.replace('sqlite:///', '', 1)
            # Resolver a ruta absoluta si es relativa
            if not os.path.isabs(path):
                path = os.path.join(os.getcwd(), path)
            if os.path.exists(path):
                return send_file(
                    path,
                    as_attachment=True,
                    download_name=f'backup_sqlite_{ts}.db'
                )
            else:
                # Si no existe el archivo (p.ej., en memoria), seguir con backup JSON
                pass
    except Exception:
        # Continuar con backup JSON si algo falla en la rama SQLite
        pass

    # Caso 2: Backup universal JSON (esquema + datos) comprimido en ZIP
    payload = {
        'database_url': db_url,
        'generated_at': _dt.utcnow().isoformat() + 'Z',
        'schema': {},
        'data': {}
    }

    try:
        inspector = inspect(db.engine)
        tablas = inspector.get_table_names()
        metadata = MetaData()

        with db.engine.connect() as conn:
            for tname in tablas:
                try:
                    # Esquema
                    cols = inspector.get_columns(tname)
                    payload['schema'][tname] = [
                        {
                            'name': c.get('name'),
                            'type': str(c.get('type')),
                            'nullable': c.get('nullable'),
                            'default': str(c.get('default')) if c.get('default') is not None else None
                        }
                        for c in cols
                    ]

                    # Datos (usando reflexión para respetar comillas de identificadores)
                    table_obj = Table(tname, metadata, autoload_with=db.engine)
                    result = conn.execute(select(table_obj))
                    rows = [dict(r._mapping) for r in result]
                    payload['data'][tname] = rows
                except Exception as te:
                    # Si una tabla falla, registrar el error y continuar
                    payload['data'][tname] = {'__error__': str(te)}

        # Empaquetar en ZIP con un archivo JSON interno
        bio = io.BytesIO()
        with zipfile.ZipFile(bio, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('backup.json', json.dumps(payload, ensure_ascii=False, default=str, indent=2))
        bio.seek(0)

        return send_file(
            bio,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'backup_adat_{ts}.zip'
        )
    except Exception as e:
        flash(f'Error al generar backup: {e}', 'danger')
        return redirect(url_for('estadisticas'))


@app.route('/admin/backup/import', methods=['POST'])
@login_required
def admin_backup_import():
    """Importa una copia completa del sistema desde un archivo de backup.
    Acepta:
    - ZIP con backup.json (esquema+datos) -> restaura limpiando e insertando.
    - .db (solo SQLite) -> reemplaza el archivo de BD (requiere reinicio de app para asegurar conexiones limpias).
    """
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))

    f = request.files.get('file')
    if not f or not getattr(f, 'filename', ''):
        flash('No se envió ningún archivo de backup', 'warning')
        return redirect(url_for('estadisticas'))

    filename = f.filename.lower()
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')

    try:
        # Caso A: ZIP con backup.json
        if filename.endswith('.zip'):
            bio = io.BytesIO(f.read())
            with zipfile.ZipFile(bio, 'r') as zf:
                if 'backup.json' not in zf.namelist():
                    flash('El ZIP no contiene backup.json', 'danger')
                    return redirect(url_for('estadisticas'))
                payload = json.loads(zf.read('backup.json').decode('utf-8'))

            data = payload.get('data', {})

            # Obtener orden de dependencia entre tablas
            inspector = inspect(db.engine)
            tablas_presentes = [t for t in inspector.get_table_names() if t in data]

            # Construir grafo de dependencias (t1 depende de t2 si t1 tiene FK a t2)
            deps = {t: set() for t in tablas_presentes}
            for t in tablas_presentes:
                try:
                    for fk in inspector.get_foreign_keys(t):
                        ref_table = fk.get('referred_table')
                        if ref_table in deps and ref_table != t:
                            deps[t].add(ref_table)
                except Exception:
                    pass

            # Orden topológico simple
            order = []
            deps_copy = {k: set(v) for k, v in deps.items()}
            while deps_copy:
                libres = [t for t, req in deps_copy.items() if not req]
                if not libres:
                    # Ciclo o restricciones no deferrables: caer en un orden alfabético para continuar
                    libres = [next(iter(deps_copy))]
                for t in libres:
                    order.append(t)
                    deps_copy.pop(t, None)
                    for v in deps_copy.values():
                        v.discard(t)

            from sqlalchemy import MetaData, Table, select
            metadata = MetaData()

            with db.engine.begin() as conn:  # transaccional
                # Desactivar FK en SQLite para facilitar borrado/carga
                try:
                    conn.exec_driver_sql('PRAGMA foreign_keys=OFF')
                except Exception:
                    pass

                # Borrar datos en orden inverso
                for t in reversed(order):
                    try:
                        table_obj = Table(t, metadata, autoload_with=db.engine)
                        conn.execute(table_obj.delete())
                    except Exception:
                        pass

                # Insertar datos en orden
                for t in order:
                    rows = data.get(t)
                    if not isinstance(rows, list):
                        # Si es un dict con error u otra forma, saltar
                        continue
                    if not rows:
                        continue
                    table_obj = Table(t, metadata, autoload_with=db.engine)
                    # Insert batch; si falla por tipos, intentar fila a fila
                    try:
                        conn.execute(table_obj.insert(), rows)
                    except Exception:
                        for row in rows:
                            try:
                                conn.execute(table_obj.insert(), row)
                            except Exception:
                                # Último recurso: ignorar fila problemática
                                pass

                # Re-activar FK en SQLite
                try:
                    conn.exec_driver_sql('PRAGMA foreign_keys=ON')
                except Exception:
                    pass

            # Intentar resetear secuencias en PostgreSQL para columnas id
            try:
                if db_url.startswith('postgresql'):
                    with db.engine.begin() as conn:
                        for t in order:
                            try:
                                # setval para columna id si existe
                                conn.exec_driver_sql(
                                    f"SELECT setval(pg_get_serial_sequence('{t}', 'id'), COALESCE((SELECT MAX(id) FROM {t}), 1), true)"
                                )
                            except Exception:
                                pass
            except Exception:
                pass

            flash('Backup importado exitosamente', 'success')
            return redirect(url_for('estadisticas'))

        # Caso B: archivo .db (solo SQLite con archivo)
        elif filename.endswith('.db'):
            if not (db_url.startswith('sqlite:///') or db_url.startswith('sqlite:////')):
                flash('Importar .db solo está soportado para SQLite', 'danger')
                return redirect(url_for('estadisticas'))
            # Resolver ruta actual de SQLite
            if db_url.startswith('sqlite:////'):
                path = '/' + db_url.replace('sqlite:////', '', 1)
            else:
                path = db_url.replace('sqlite:///', '', 1)
            if not os.path.isabs(path):
                path = os.path.join(os.getcwd(), path)

            # Guardar a un archivo temporal y luego reemplazar
            tmp_dir = os.path.join('instance', 'backups')
            os.makedirs(tmp_dir, exist_ok=True)
            tmp_path = os.path.join(tmp_dir, 'uploaded_backup.db')
            f.save(tmp_path)

            try:
                # Cerrar conexiones antes de reemplazar (best-effort)
                db.session.close()
            except Exception:
                pass

            import shutil
            shutil.copyfile(tmp_path, path)

            flash('Base de datos SQLite importada. Reinicia la aplicación para aplicar completamente los cambios.', 'success')
            return redirect(url_for('estadisticas'))

        else:
            flash('Formato no soportado. Usa un .zip (backup.json) o .db (SQLite).', 'warning')
            return redirect(url_for('estadisticas'))

    except Exception as e:
        db.session.rollback()
        flash(f'Error al importar backup: {e}', 'danger')
        return redirect(url_for('estadisticas'))


# ==================== CONFIGURACIÓN DEL SISTEMA ====================

# Modelo para configuraciones
class ConfiguracionSistema(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    clave = db.Column(db.String(100), unique=True, nullable=False)
    valor = db.Column(db.String(5000), nullable=False)
    descripcion = db.Column(db.String(500))


# Modelo para actividades generadas
class ActividadGenerada(db.Model):
    __tablename__ = 'actividad_generada'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    nombre_profesor = db.Column(db.String(200), nullable=False)
    grado = db.Column(db.String(50), nullable=False)
    asignatura = db.Column(db.String(200), nullable=False)
    tematica = db.Column(db.String(500), nullable=False)
    cantidad_estudiantes = db.Column(db.Integer, nullable=False)
    tipo_actividad = db.Column(db.String(50), nullable=False)
    tiempo = db.Column(db.Integer, nullable=False)
    recursos = db.Column(db.String(1000), nullable=False)
    contenido = db.Column(db.Text, nullable=False)
    fecha_creacion = db.Column(db.DateTime, nullable=False, default=datetime.now)
    fecha_modificacion = db.Column(db.DateTime, nullable=False, default=datetime.now, onupdate=datetime.now)
    
    # Relación con usuario
    usuario = db.relationship('User', backref=db.backref('actividades_generadas', lazy='dynamic'))
    
    def __repr__(self):
        return f'<ActividadGenerada {self.asignatura} - {self.tematica}>'


# ====== AUTO-CREACIÓN SEGURA DE LA TABLA ACTIVIDAD_GENERADA ======
# No usa el ambiente/venv. En el arranque valida y crea solo esta tabla si falta.
try:
    with app.app_context():
        inspector = inspect(db.engine)
        if not inspector.has_table('actividad_generada'):
            # Crear únicamente la tabla del modelo ActividadGenerada
            ActividadGenerada.__table__.create(bind=db.engine, checkfirst=True)
            # Crear índice para user_id (idempotente)
            db.session.execute(text('CREATE INDEX IF NOT EXISTS idx_actividad_user_id ON actividad_generada(user_id)'))
            db.session.commit()
            print('✅ Tabla actividad_generada creada automáticamente')
        else:
            # Asegurar índice (no falla si ya existe)
            db.session.execute(text('CREATE INDEX IF NOT EXISTS idx_actividad_user_id ON actividad_generada(user_id)'))
            db.session.commit()
except Exception as e:
    # No bloquear arranque si hay un problema; registrar y continuar
    try:
        db.session.rollback()
    except Exception:
        pass
    print(f'⚠️  No se pudo verificar/crear la tabla actividad_generada: {e}')


# ====== AJUSTE AUTOMÁTICO DE TIPO PARA ENUNCIADOS LARGOS (QUESTION.statement) ======
try:
    with app.app_context():
        # Solo aplicar a motores que soportan ALTER COLUMN sencillo (PostgreSQL principalmente).
        dialect = db.engine.name  # 'postgresql', 'sqlite', etc.
        if dialect == 'postgresql':
            insp = inspect(db.engine)
            tablas = insp.get_table_names()
            for tname in ('question', 'quiz1_question'):
                if tname in tablas:
                    cols = insp.get_columns(tname)
                    st = next((c for c in cols if c.get('name') == 'statement'), None)
                    # Si es VARCHAR o tiene length limitado, migrar a TEXT
                    if st and (str(st.get('type')).upper().startswith('VARCHAR') or 'STRING' in str(st.get('type')).upper()):
                        try:
                            db.session.execute(text(f'ALTER TABLE {tname} ALTER COLUMN statement TYPE TEXT'))
                            db.session.commit()
                            print(f'✅ Columna {tname}.statement convertida a TEXT (soporta enunciados largos)')
                        except Exception as e2:
                            db.session.rollback()
                            print(f'⚠️  No se pudo convertir {tname}.statement a TEXT: {e2}')
        # Para SQLite, no es necesario; maneja long text sin truncado práctico.
except Exception as e:
    try:
        db.session.rollback()
    except Exception:
        pass
    print(f'⚠️  Error en ajuste automático de statement: {e}')

# ====== AGREGAR COLUMNA CEDULA SI NO EXISTE ======
try:
    with app.app_context():
        insp = inspect(db.engine)
        if 'user' in insp.get_table_names():
            cols = insp.get_columns('user')
            col_names = [c.get('name') for c in cols]
            if 'cedula' not in col_names:
                try:
                    db.session.execute(text('ALTER TABLE "user" ADD COLUMN cedula VARCHAR(20)'))
                    db.session.commit()
                    print('✅ Columna cedula agregada a la tabla user')
                except Exception as e2:
                    db.session.rollback()
                    print(f'⚠️  No se pudo agregar columna cedula: {e2}')
except Exception as e:
    try:
        db.session.rollback()
    except Exception:
        pass
    print(f'⚠️  Error al verificar/agregar columna cedula: {e}')


@app.route('/admin/configuracion', methods=['GET', 'POST'])
@login_required
def configuracion_sistema():
    """Panel de configuración del sistema"""
    if current_user.rol != 'admin':
        flash('No tienes permisos para acceder a esta página', 'danger')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        try:
            # Actualizar configuraciones
            for key in request.form:
                if key.startswith('config_'):
                    clave = key.replace('config_', '')
                    valor = request.form.get(key)
                    
                    config = ConfiguracionSistema.query.filter_by(clave=clave).first()
                    if config:
                        config.valor = valor
                    else:
                        config = ConfiguracionSistema(clave=clave, valor=valor)
                        db.session.add(config)
            
            db.session.commit()
            flash('Configuración actualizada exitosamente', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar configuración: {str(e)}', 'danger')
    
    # Obtener o crear configuraciones por defecto
    configs = {
        'quiz1_activo': ConfiguracionSistema.query.filter_by(clave='quiz1_activo').first() or 
                        ConfiguracionSistema(clave='quiz1_activo', valor='true', descripcion='Activar/Desactivar Quiz 1 (ADAT)'),
        'quiz2_activo': ConfiguracionSistema.query.filter_by(clave='quiz2_activo').first() or 
                        ConfiguracionSistema(clave='quiz2_activo', valor='true', descripcion='Activar/Desactivar Quiz 2 (Estilos)'),
        'quiz3_activo': ConfiguracionSistema.query.filter_by(clave='quiz3_activo').first() or 
                        ConfiguracionSistema(clave='quiz3_activo', valor='true', descripcion='Activar/Desactivar Quiz 3 (Jugador)'),
        'quiz4_activo': ConfiguracionSistema.query.filter_by(clave='quiz4_activo').first() or 
                        ConfiguracionSistema(clave='quiz4_activo', valor='true', descripcion='Activar/Desactivar Quiz 4 (Pensamiento)'),
        'generador_actividades_activo': ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first() or 
                         ConfiguracionSistema(clave='generador_actividades_activo', valor='true', descripcion='Activar/Desactivar Generador de Actividades y Mis Actividades'),
        'resultados_activo': ConfiguracionSistema.query.filter_by(clave='resultados_activo').first() or 
                         ConfiguracionSistema(clave='resultados_activo', valor='true', descripcion='Activar/Desactivar vista de Resultados del usuario'),
        'perfil_activo': ConfiguracionSistema.query.filter_by(clave='perfil_activo').first() or 
                         ConfiguracionSistema(clave='perfil_activo', valor='true', descripcion='Activar/Desactivar vista de Perfil del usuario'),
        'mensaje_bienvenida': ConfiguracionSistema.query.filter_by(clave='mensaje_bienvenida').first() or 
                              ConfiguracionSistema(clave='mensaje_bienvenida', valor='Bienvenido a la plataforma ADAT', descripcion='Mensaje de bienvenida'),
        'permitir_registro': ConfiguracionSistema.query.filter_by(clave='permitir_registro').first() or 
                             ConfiguracionSistema(clave='permitir_registro', valor='true', descripcion='Permitir nuevos registros'),
        'tiempo_quiz1': ConfiguracionSistema.query.filter_by(clave='tiempo_quiz1').first() or 
                        ConfiguracionSistema(clave='tiempo_quiz1', valor='30', descripcion='Tiempo límite Quiz 1 (minutos)'),
        'tiempo_quiz4': ConfiguracionSistema.query.filter_by(clave='tiempo_quiz4').first() or 
                        ConfiguracionSistema(clave='tiempo_quiz4', valor='45', descripcion='Tiempo límite Quiz 4 (minutos)'),
    }
    
    # Guardar configs nuevas si no existen
    for config in configs.values():
        if not config.id:
            db.session.add(config)
    try:
        db.session.commit()
    except:
        db.session.rollback()
    
    return render_template('admin_configuracion.html', configs=configs)


# Exponer configuración al contexto de plantillas
@app.context_processor
def inject_config_flags():
    """Inyecta helpers y banderas de configuración en todas las plantillas"""
    def cfg(clave, default=None):
        try:
            c = ConfiguracionSistema.query.filter_by(clave=clave).first()
            return c.valor if c else default
        except Exception:
            return default

    # Función para convertir fecha UTC a hora de Bogotá (UTC-5)
    def format_datetime_bogota(dt):
        if not dt:
            return 'Nunca'
        # Bogotá está en UTC-5
        bogota_time = dt - timedelta(hours=5)

        return bogota_time.strftime('%d/%m/%Y %H:%M')

    # Bandera para generador de actividades (por defecto True)
    try:
        c = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        generador_activo = (c.valor == 'true') if c else True
    except Exception:
        generador_activo = True

    return dict(cfg=cfg, format_datetime_bogota=format_datetime_bogota, generador_actividades_activo=generador_activo)


@app.route('/register_quiz1_question', methods=['GET', 'POST'])
def register_quiz1_question():
    form = QuizForm()
    if form.validate_on_submit():
        question = Question(
            statement=form.statement.data,
            option_a=form.option_a.data,
            option_b=form.option_b.data,
            option_c=form.option_c.data,
            option_d=form.option_d.data,
            correct_answer=form.correct_answer.data,
            label=form.label.data,
            percentage=form.percentage.data,
            image_url=form.image_url.data if form.image_url.data else ''
        )
        db.session.add(question)
        db.session.commit()
        flash('Pregunta registrada con éxito', 'success')
        return redirect(url_for('register_quiz1_question'))
    return render_template('register_quiz1_question.html', form=form)


@app.route('/results', methods=['GET'])
def results():
    # Aquí deberías calcular los resultados basados en las respuestas del usuario
    results_data = [
        {'label': 'Etiqueta 1', 'percentage': 75},
        {'label': 'Etiqueta 2', 'percentage': 50},
        {'label': 'Etiqueta 3', 'percentage': 90},
    ]

    user_responses = {
        "¿Cuál es tu color favorito?": "Azul",
        "¿Te gusta la programación?": "Sí",
    }

    return render_template('results.html', results_data=results_data, user_responses=user_responses)


@app.route('/quiz1', methods=['GET', 'POST'])
@login_required
def quiz1():
    # Flag de acceso a ADAT
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='quiz1_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El test ADAT está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    resultado = ResultadoQuiz.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "success")
        return redirect(url_for('quiz_results', result_id=resultado.id))

    page = request.args.get('page', 1, type=int)  # Obtener el número de página
    questions_per_page = 15  # Número de preguntas por página
    # Obtén todas las preguntas del quiz 1
    questions = Question.query.order_by(Question.id).all()
    total_questions = len(questions)
    
    # Calcular las preguntas para la página actual
    start = (page - 1) * questions_per_page
    end = start + questions_per_page
    questions_to_display = questions[start:end]

    if request.method == 'POST':
        # Aquí puedes procesar las respuestas del formulario
        # Por ejemplo, puedes guardar los resultados en la base de datos
        user_answers = request.form  # Recoge las respuestas del usuario

        # Procesa las respuestas, verifica y guarda los resultados...

        return redirect(url_for('results'))  # Redirige a la página de resultados

    return render_template('quiz1.html', questions=questions_to_display, page=page, total_questions=total_questions, questions_per_page=questions_per_page)

@app.route('/submit_quiz', methods=['POST'])
@login_required
def submit_quiz():
    user_responses = request.form.to_dict()
    resultado = ResultadoQuiz.query.filter_by(user_id=current_user.id).first()
    print("Respuestas del usuario:", user_responses)

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "warning")
        return redirect(url_for('quiz_results', result_id=resultado.id))

    correct_count = 0
    incorrect_count = 0
    etiqueta_scores = {
        'Abstracción': 0,
        'Descomposición': 0,
        'Pensamiento Algoritmico': 0,
    }

    for question_id, user_answer in user_responses.items():
        # Aquí debes obtener la pregunta desde la base de datos
        question = Question.query.get(int(question_id))
        if question is not None:
            # Comprobar si la respuesta del usuario es correcta
            if question.correct_answer == user_answer:
                correct_count += 1
                etiqueta_scores[question.label] = etiqueta_scores.get(question.label, 0) + question.percentage
            else:
                incorrect_count += 1
        else:
            print(f"Pregunta con ID {question_id} no encontrada en la base de datos.")

    # Serializa etiqueta_scores para pasarlo como argumento en la URL
    etiqueta_scores_json = json.dumps({k: float(v) for k, v in etiqueta_scores.items()})
    #etiqueta_scores_json = json.dumps({k: float(v) if v is not None else 0 for k, v in etiqueta_scores.items()})


    # Guardar en la base de datos
    resultado = ResultadoQuiz(
        user_id=current_user.id,
        abstraccion=etiqueta_scores.get('Abstracción', 0),
        descomposicion=etiqueta_scores.get('Descomposición', 0),
        pensamiento_algoritmico=etiqueta_scores.get('Pensamiento Algorítmico', 0),
        respuestas_correctas=correct_count,
        respuestas_incorrectas=incorrect_count,
    )
    db.session.add(resultado)
    db.session.commit()


    return redirect(url_for('quiz_results', result_id=resultado.id))

@app.route('/submit_quiz4', methods=['POST'])
@login_required
def submit_quiz4():
    user_responses = request.form.to_dict()
    
    # Verificar si el usuario ya tiene un resultado - hacer esto de manera más robusta
    resultado_existente = ResultadoQuizCuatro.query.filter_by(user_id=current_user.id).first()
    if resultado_existente:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "warning")
        return redirect(url_for('quiz4_results', result_id=resultado_existente.id))

    correct_count = 0
    incorrect_count = 0
    score = 0

    for question_id, user_answer in user_responses.items():
        # Aquí debes obtener la pregunta desde la base de datos
        question = QuestionCuatro.query.get(int(question_id))
        if question is not None:
            # Comprobar si la respuesta del usuario es correcta
            if question.correct_answer == user_answer:
                correct_count += 1
                score += question.percentage if question.percentage else 1
            else:
                incorrect_count += 1
        else:
            print(f"Pregunta con ID {question_id} no encontrada en la base de datos.")

    # Guardar en la base de datos del Quiz 4
    try:
        resultado = ResultadoQuizCuatro(
            user_id=current_user.id,
            score=int(score),
            respuestas_correctas=correct_count,
            respuestas_incorrectas=incorrect_count,
        )
        db.session.add(resultado)
        db.session.commit()
        return redirect(url_for('quiz4_results', result_id=resultado.id))
    except IntegrityError as e:
        db.session.rollback()
        # Verificar nuevamente si ya existe un resultado (por si fue creado por otra transacción)
        resultado_existente = ResultadoQuizCuatro.query.filter_by(user_id=current_user.id).first()
        if resultado_existente:
            flash("Ya has completado este cuestionario. Aquí están tus resultados.", "warning")
            return redirect(url_for('quiz4_results', result_id=resultado_existente.id))
        else:
            # Si no existe pero aún así hay error de integridad, es un problema diferente
            flash("Error al guardar los resultados. Por favor, intenta nuevamente.", "danger")
            return redirect(url_for('quiz4'))
    except Exception as e:
        db.session.rollback()
        print(f"Error inesperado al guardar resultado del quiz 4: {e}")
        flash("Error al guardar los resultados. Por favor, intenta nuevamente.", "danger")
        return redirect(url_for('quiz4'))


@app.route('/quiz_results/<int:result_id>')
def quiz_results(result_id):
    resultado = ResultadoQuiz.query.get_or_404(result_id)

    correct_count = request.args.get('correct_count', type=int)
    incorrect_count = request.args.get('incorrect_count', type=int)
    etiqueta_scores = {
        "Abstracción": resultado.abstraccion,
        "Descomposición": resultado.descomposicion,
        "Pensamiento Algoritmico": resultado.pensamiento_algoritmico
    }
    
    # Convertir el string de etiqueta_scores a un diccionario
    etiqueta_scores_json = json.dumps({k: float(v) for k, v in etiqueta_scores.items()})

    # Calcular el porcentaje total para cada etiqueta
    total_percentage = sum(etiqueta_scores.values())
    if total_percentage > 0:
        for label in etiqueta_scores:
            etiqueta_scores[label] = round((etiqueta_scores[label] / total_percentage) * 100, 2)

    # Diccionario con descripciones de cada etiqueta
    etiqueta_descriptions = {
        'Abstracción': 'Entender el problema.',
        'Descomposición': 'Dividir problemas complejos en partes más simples.',
        'Pensamiento Algoritmico': 'Proporcionar soluciones paso a paso a un problema.',
    }

    
    return render_template('quiz_results.html', 
                           resultado=resultado, 
                           etiqueta_scores=etiqueta_scores, etiqueta_descriptions=etiqueta_descriptions )



@app.route('/register_quiz2_question', methods=['GET', 'POST'])
def register_quiz2_question():
    form = QuizFormDos()
    if form.validate_on_submit():
        question = QuestionDos(
            statement=form.statement.data,
            option_a=form.option_a.data,
            option_b=form.option_b.data,
        )
        db.session.add(question)
        db.session.commit()
        flash('Pregunta registrada con éxito', 'success')
        return redirect(url_for('register_quiz2_question'))
    return render_template('register_quiz2_question.html', form=form)


@app.route('/register_quiz4_question', methods=['GET', 'POST'])
def register_quiz4_question():
    form = QuizFormCuatro()
    if form.validate_on_submit():
        # Manejar upload o URL
        uploaded = request.files.get('image_file')
        image_url_final = form.image_url.data if form.image_url.data else ''

        question = QuestionCuatro(
            statement=form.statement.data,
            option_a=form.option_a.data,
            option_b=form.option_b.data,
            option_c=form.option_c.data,
            option_d=form.option_d.data,
            correct_answer=form.correct_answer.data,
            label=form.label.data or 'CTt',
            percentage=form.percentage.data or 1,
            image_url=image_url_final
        )
        db.session.add(question)
        db.session.commit()
        flash('Pregunta registrada con éxito', 'success')
        return redirect(url_for('register_quiz4_question'))
    return render_template('register_quiz4_question.html', form=form)

@app.route('/quiz2', methods=['GET', 'POST'])
@login_required
def quiz2():
    # Flag de acceso a Estilos de aprendizaje
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='quiz2_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El test de Estilos de Aprendizaje está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    resultado = ResultadoQuizDos.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "success")
        return redirect(url_for('quiz_resultsDos', result_id=resultado.id))

    page = request.args.get('page', 1, type=int)  # Obtener el número de página
    questions_per_page = 45  # Número de preguntas por página
    # Obtén todas las preguntas del quiz 2
    questions = QuestionDos.query.order_by(QuestionDos.id).all()
    total_questions = len(questions)
    
    # Calcular las preguntas para la página actual
    start = (page - 1) * questions_per_page
    end = start + questions_per_page
    questions_to_display = questions[start:end]

    if request.method == 'POST':
        # Aquí puedes procesar las respuestas del formulario
        # Por ejemplo, puedes guardar los resultados en la base de datos
        user_answers = request.form  # Recoge las respuestas del usuario

        # Procesa las respuestas, verifica y guarda los resultados...

        return redirect(url_for('results2'))  # Redirige a la página de resultados

    return render_template('quiz2.html', questions=questions_to_display, page=page, total_questions=total_questions, questions_per_page=questions_per_page)


@app.route('/submit_quizDos', methods=['POST'])
@login_required
def submit_quizDos():
    user_responses = request.form.to_dict()
    resultado = ResultadoQuizDos.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "warning")
        return redirect(url_for('quiz_resultsDos', result_id=resultado.id))

    print("Respuestas del usuario:", user_responses)
    etiquetas = {
        'Sensorial-Intuitivo': {'A': 0, 'B': 0},
        'Visual-Verbal': {'A': 0, 'B': 0},
        'Activo-Reflexivo': {'A': 0, 'B': 0},
        'Secuencial-Global': {'A': 0, 'B': 0},
    }

    pregunta_a_etiqueta = {
        1: 'Activo-Reflexivo',
        2: 'Sensorial-Intuitivo',
        3: 'Visual-Verbal',
        4: 'Secuencial-Global',
        5: 'Activo-Reflexivo',
        6: 'Sensorial-Intuitivo',
        7: 'Visual-Verbal',
        8: 'Secuencial-Global',
        9: 'Activo-Reflexivo',
        10: 'Sensorial-Intuitivo',
        11: 'Visual-Verbal',
        12: 'Secuencial-Global',
        13: 'Activo-Reflexivo',
        14: 'Sensorial-Intuitivo',
        15: 'Visual-Verbal',
        16: 'Secuencial-Global',
        17: 'Activo-Reflexivo',
        18: 'Sensorial-Intuitivo',
        19: 'Visual-Verbal',
        20: 'Secuencial-Global',
        21: 'Activo-Reflexivo',
        22: 'Sensorial-Intuitivo',
        23: 'Visual-Verbal',
        24: 'Secuencial-Global',
        25: 'Activo-Reflexivo',
        26: 'Sensorial-Intuitivo',
        27: 'Visual-Verbal',
        28: 'Secuencial-Global',
        29: 'Activo-Reflexivo',
        30: 'Sensorial-Intuitivo',
        31: 'Visual-Verbal',
        32: 'Secuencial-Global',
        33: 'Activo-Reflexivo',
        34: 'Sensorial-Intuitivo',
        35: 'Visual-Verbal',
        36: 'Secuencial-Global',
        37: 'Activo-Reflexivo',
        38: 'Sensorial-Intuitivo',
        39: 'Visual-Verbal',
        40: 'Secuencial-Global',
        41: 'Activo-Reflexivo',
        42: 'Sensorial-Intuitivo',
        43: 'Visual-Verbal',
        44: 'Secuencial-Global',

    }


    for question_id_str, user_answer in user_responses.items():
        question_id = int(question_id_str)
        etiqueta = pregunta_a_etiqueta.get(question_id)

        if etiqueta:
            etiquetas[etiqueta][user_answer] += 1  # Simplificado

    # Calcular el resultado final para cada etiqueta
    etiqueta_resultados = {}
    for etiqueta, conteo in etiquetas.items():
        diferencia = conteo['A'] - conteo['B']

        # Determinar el estado y valor basado en la diferencia
        if diferencia >= 9:
            valor = 'Fuerte'
        elif 5 <= diferencia <= 8:
            valor = 'Moderado'
        else:
            valor = 'Apropiado'

        if conteo['A'] > conteo['B']:
            estado = etiqueta.split('-')[0]  # Ejemplo: 'Sensorial'
            total = conteo['A']
        else:
            estado = etiqueta.split('-')[1]  # Ejemplo: 'Intuitivo'
            total = conteo['B']

        etiqueta_resultados[etiqueta] = {
            'estado': estado,
            'valor': valor,
            'total': total
        }

    print("Resultados de etiquetas:", etiqueta_resultados)

    resultado = ResultadoQuizDos(
        user_id=current_user.id,
        sensorial_intuitivo=etiqueta_resultados['Sensorial-Intuitivo'],
        visual_verbal=etiqueta_resultados['Visual-Verbal'],
        activo_reflexivo=etiqueta_resultados['Activo-Reflexivo'],
        secuencial_global=etiqueta_resultados['Secuencial-Global'],
    )

    db.session.add(resultado)
    db.session.commit()

    # Redirigir a la página de resultados con los resultados como argumento
    return redirect(url_for('quiz_resultsDos', result_id=resultado.id))


@app.route('/quiz_resultsDos/<int:result_id>')
def quiz_resultsDos(result_id):
    result = ResultadoQuizDos.query.get_or_404(result_id)

    etiqueta_resultados = {
        'Sensitivo-Intuitivo': result.sensorial_intuitivo,
        'Visual-Verbal': result.visual_verbal,
        'Activo-Reflexivo': result.activo_reflexivo,
        'Secuencial-Global': result.secuencial_global,
    }

    # Diccionario con descripciones de cada etiqueta
    etiqueta_descriptions = {
        'Sensitivo-Intuitivo': 'Como perciben la información. Los sensitivos son concretos, prácticos, orientados hacia hechos y procedimientos; les gusta resolver problemas siguiendo procedimientos muy bien establecidos. Los intuitivos son conceptuales, innovadores, orientados hacia las teorías y los significados; les gusta innovar y odian la repetición.',
        'Visual-Verbal': 'En que canal perciben mejor la información. Los visuales prefieren representaciones visuales, diagramas de flujo, imagenes, etc. Los verbales prefieren obtener la información en forma escrita o hablada; recuerdan mejor lo que leen o lo que oyen.',
        'Activo-Reflexivo': 'Como procesan la información. Los activos tienden a retener y comprender mejor nueva información cuando hacen algo activo con ella (discutiéndola, aplicándola, explicándosela a otros). Los reflexivos tienden a retener y comprender nueva información pensando y reflexionando sobre ella, prefieren aprender meditando, pensando y trabajando solos.',
        'Secuencial-Global': 'Como progresa el estudiante en su aprendizaje. Los secuenciales aprenden en pequeños pasos incrementales cuando el siguiente paso está siempre lógicamente relacionado con el anterior; ordenados y lineales. Los globales aprenden grandes saltos, aprendiendo nuevo material casi al azar y visualizando la totalidad; pueden resolver problemas complejos. Pueden tener dificultades, sin embargo, en explicar cómo lo hicieron.',

        }
    etiqueta_resultados = dict(sorted(etiqueta_resultados.items(), key=lambda x: x[1]['total'], reverse=True))

    return render_template('quiz_resultsDos.html', etiqueta_resultados=etiqueta_resultados, etiqueta_descriptions=etiqueta_descriptions)


@app.route('/register_quiz3_question', methods=['GET', 'POST'])
def register_quiz3_question():
    form = QuizFormTres()
    if form.validate_on_submit():
        question = QuestionTres(
            statement=form.statement.data,
            option_a=form.option_a.data,
            option_b=form.option_b.data,
            option_c=form.option_c.data,
            option_d=form.option_d.data,
            option_e=form.option_e.data,

        )
        db.session.add(question)
        db.session.commit()
        flash('Pregunta registrada con éxito', 'success')
        return redirect(url_for('register_quiz3_question'))
    return render_template('register_quiz3_question.html', form=form)


@app.route('/quiz3', methods=['GET', 'POST'])
@login_required
def quiz3():
    # Flag de acceso a Tipo de Jugador
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='quiz3_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El test de Tipo de Jugador está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    resultado = ResultadoQuizTres.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "success")
        return redirect(url_for('quiz_resultsTres', result_id=resultado.id))

    page = request.args.get('page', 1, type=int)  # Obtener el número de página
    questions_per_page = 12  # Número de preguntas por página
    # Obtén todas las preguntas del quiz 3
    questions = QuestionTres.query.order_by(QuestionTres.id).all()
    total_questions = len(questions)
    
    # Calcular las preguntas para la página actual
    start = (page - 1) * questions_per_page
    end = start + questions_per_page
    questions_to_display = questions[start:end]

    if request.method == 'POST':
        # Aquí puedes procesar las respuestas del formulario
        # Por ejemplo, puedes guardar los resultados en la base de datos
        user_answers = request.form.to_dict()  # Recoge las respuestas del usuario
        print("RespuestasX", user_answers)
        # Procesa las respuestas, verifica y guarda los resultados...


        return redirect(url_for('results3'))  # Redirige a la página de resultados

    return render_template('quiz3.html', questions=questions_to_display, page=page, total_questions=total_questions, questions_per_page=questions_per_page)


@app.route('/submit_quizTres', methods=['POST'])
@login_required
def submit_quizTres():
    user_answers = request.form.to_dict()
    resultado = ResultadoQuizTres.query.filter_by(user_id=current_user.id).first()

    if resultado:
        flash("Ya has completado este cuestionario. Aquí están tus resultados.", "warning")
        return redirect(url_for('quiz_resultsTres', result_id=resultado.id))


    etiquetas = {
        'Filántropo': 0,
        'Socializador': 0,
        'Triunfador': 0,
        'Jugador': 0,
        'Espíritu Libre': 0,
        'Disruptor': 0
        }
    total = 0

     # Asociación de preguntas con etiquetas
    pregunta_a_etiqueta = {
        1: 'Filántropo',
        2: 'Filántropo',
        3: 'Socializador',
        4: 'Socializador',
        5: 'Triunfador',
        6: 'Triunfador',
        7: 'Jugador',
        8: 'Jugador',
        9: 'Espíritu Libre',
        10: 'Espíritu Libre',
        11: 'Disruptor',
        12: 'Disruptor'
    }

    for question_id_str, user_answer in user_answers.items():
        # Aquí debes obtener la pregunta desde la base de datos
        #question = QuestionTres.query.get(int(question_id))
        question_id = (int(question_id_str))
        valor = 0

        if user_answer == 'A':
            valor = 5
        elif user_answer == "B":
            valor = 4
        elif user_answer == "C":
            valor = 3
        elif user_answer == "D":
            valor = 2
        elif user_answer == "E":
            valor = 1
        
        print(f"Pregunta ID: {question_id}, Valor: {valor}")  # Agregar esta línea para ver el valor asignado

        etiqueta = pregunta_a_etiqueta.get(question_id)
        print(f"Pregunta ID: {question_id}, Valor: {valor}, Etiqueta: {etiqueta}")

        if etiqueta:
            etiquetas[etiqueta] += valor
            
        total += valor
        

    # Calcular el porcentaje para cada etiqueta
    etiqueta_percentages = {etiqueta: round((score / total) * 100, 2) if total > 0 else 0 
                            for etiqueta, score in etiquetas.items()}

    # Guardar los resultados en la base de datos
    # (aquí deberías añadir la lógica para guardar en la base de datos si es necesario)

    #return render_template('quiz_resultsTres.html', etiqueta_percentages=etiqueta_percentages)
     # Guardar en la base de datos
    resultado = ResultadoQuizTres(
        user_id=current_user.id,
        filantropo=etiqueta_percentages['Filántropo'],
        socializador=etiqueta_percentages['Socializador'],
        triunfador=etiqueta_percentages['Triunfador'],
        jugador=etiqueta_percentages['Jugador'],
        espiritu_libre=etiqueta_percentages['Espíritu Libre'],
        disruptor=etiqueta_percentages['Disruptor']
    )

    db.session.add(resultado)
    db.session.commit()

    return redirect(url_for('quiz_resultsTres', result_id=resultado.id))



@app.route('/quiz_resultsTres/<int:result_id>')
def quiz_resultsTres(result_id):
    result = ResultadoQuizTres.query.get_or_404(result_id)
    
    etiqueta_percentages = {
        'Filántropo': result.filantropo,
        'Socializador': result.socializador,
        'Triunfador': result.triunfador,
        'Jugador': result.jugador,
        'Espíritu Libre': result.espiritu_libre,
        'Disruptor': result.disruptor
    }

    # Diccionario con descripciones de cada etiqueta
    etiqueta_descriptions = {
        'Filántropo': 'Impulsado por un propósito en particular. No necesitan ninguna recompensa, con obtener la sensación de aportar un valor social ya es suficiente.',
        'Socializador': 'Impulsado por las relaciones interpersonales y motivados por la interacción con otros jugadores.',
        'Triunfador': 'Motivados por la competencia y la maestría. Busca desafíos, metas y recompensas para demostrar su habilidad y esfuerzo.',
        'Jugador': 'Impulsado por las recompensas. Se involucra a través de la competencia y disfruta ganar en distintos escenarios.',
        'Espíritu Libre': 'Impulsado por la autonomía y la libertad. Valora la creatividad y la exploración.',
        'Disruptor': 'Impulsado por el cambio. Le gusta desafiar lo establecido, experimentar y generar cambios innovadores.',
    }
    # Recupera los porcentajes de las etiquetas desde los argumentos de la URL
    #etiqueta_percentages_str = request.args.get('etiqueta_percentages')
    
    # Ordenar las etiquetas de mayor a menor porcentaje
    etiqueta_percentages = dict(sorted(etiqueta_percentages.items(), key=lambda x: x[1], reverse=True))

    return render_template('quiz_resultsTres.html', etiqueta_percentages=etiqueta_percentages, etiqueta_descriptions=etiqueta_descriptions)


@app.route('/get_form_data', methods=['GET'])
def get_form_data():
    grados = Grado.query.all()
    # Modelos eliminados: Asignatura, Tematica, Habilidad. Retornar listas vacías.
    form_data = {
        'grados': [{'id': g.id, 'nombre': g.nombre} for g in grados],
        'asignaturas': [],
        'tematicas': [],
        'habilidades': []
    }
    return jsonify(form_data)


# Ruta para el autocompletar de recursos
@app.route('/autocomplete_recursos', methods=['GET'])
def autocomplete_recursos():
    # Modelo Recurso eliminado. Retornar lista vacía.
    return jsonify([])

# Función para generar actividades dinámicas usando ChatGPT
def generate_activity(grade, subject, topic, skill, students, time, resources):
    prompt = f"""
    Genera una actividad de clase para {students} estudiantes de {grade} grado sobre la temática de {topic} en la asignatura {subject}.
    La actividad debe desarrollar la habilidad de {skill} y durar {time} minutos. Además, se deben usar los siguientes recursos: {resources}.
    Proporcióname una actividad creativa que un docente pueda realizar en clase.
    """
    
    # Llamada a la API de OpenAI para generar el texto
    if not openai.api_key:
        return "Configura OPENAI_API_KEY para generar actividades automáticamente."
    try:
        response = openai.completions.create(
            model="gpt-3.5-turbo",
            prompt=prompt,
            max_tokens=200
        )
        # SDKs recientes devuelven dict-like; maneja ambas formas
        if hasattr(response, 'choices'):
            choice = response.choices[0]
            text = getattr(choice, 'text', None) or choice.get('text', '')
            return text.strip()
        return "No se recibió un texto de respuesta de OpenAI."
    except Exception as e:
        return f"Error al generar actividades: {e}"


@app.route('/crear_actividad', methods=['GET', 'POST'])
def crear_actividad():
    grados = Grado.query.all()  # Consulta todos los grados
    # Modelos eliminados: Asignatura, Habilidad. Usar listas vacías.
    asignaturas = []
    habilidades = []

    if request.method == 'POST':
        grado = request.form.get('grado')
        asignatura = request.form.get('asignatura')
        tematica = request.form.get('tematica')
        habilidad = request.form.getlist('habilidad')
        cantidad_estudiantes = request.form.get('cantidad_estudiantes')
        cantidad_tiempo = request.form.get('cantidad_tiempo')
        recursos = request.form.get('recursos')

        # Genera actividades utilizando el modelo GPT de OpenAI
        prompt = f"Genera actividades para el grado {grado}, en la asignatura {asignatura}, con la temática {tematica}, enfocadas en las habilidades {', '.join(habilidad)} para {cantidad_estudiantes} estudiantes en un tiempo de {cantidad_tiempo} minutos. Recursos disponibles: {recursos}."

        try:
            # Uso de la nueva API de OpenAI
            response = openai.completions.create(
                model="gpt-3.5-turbo",  # Puedes usar gpt-4 si tienes acceso
                prompt=prompt,
                max_tokens=150,
                n=1,
                stop=None,
                temperature=0.7
            )

            actividades_generadas = response['choices'][0]['text'].strip()
            return render_template('select_activities.html', actividades=actividades_generadas)

        except Exception as e:
            return f"Error al generar actividades: {e}"

    grados = Grado.query.all()
    # Modelos eliminados: Asignatura, Habilidad. Usar listas vacías.
    asignaturas = []
    habilidades = []
    return render_template('select_activities.html', grados=grados, asignaturas=asignaturas, habilidades=habilidades)


@app.route('/select_activities', methods=['GET', 'POST'])
@login_required
def select_activities():
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El Generador de Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        # Obtener datos del formulario
        nombre_profesor = request.form.get('nombre_profesor')
        grado = request.form.get('grado')
        asignatura = request.form.get('asignatura')
        tematica = request.form.get('tematica')
        cantidad_estudiantes = request.form.get('cantidad_estudiantes')
        tipo_actividad = request.form.get('tipo_actividad')  # individual o colaborativo
        tiempo = request.form.get('tiempo')
        recursos = request.form.get('recursos')
        competencias = request.form.getlist('competencias')  # lista de competencias seleccionadas

        # Generar actividad usando OpenAI
        try:
            # Preparar cadena legible de competencias
            competencias_txt = ", ".join(competencias) if competencias else "(no especificado)"

            prompt = f"""
Eres un experto en pedagogía y diseño de actividades educativas. Crea un taller educativo completo con la siguiente información:

- Asignatura: {asignatura}
- Grado: {grado}
- Temática: {tematica}
- Profesor: {nombre_profesor}
- Cantidad de estudiantes: {cantidad_estudiantes}
- Tipo de actividad: {tipo_actividad}
- Tiempo disponible: {tiempo} minutos
- Recursos disponibles: {recursos}
 - Competencias a desarrollar: {competencias_txt}

Genera un taller estructurado que incluya:
1. Título creativo y llamativo
2. Objetivos de aprendizaje (3-4 objetivos claros)
3. Introducción motivadora
4. Desarrollo de la actividad (paso a paso)
5. Actividades prácticas específicas
6. Evaluación o cierre
7. Recursos necesarios

El taller debe ser apropiado para el nivel educativo, aprovechar los recursos disponibles, y adaptarse al tiempo y tipo de trabajo ({tipo_actividad}).
Usa un lenguaje claro y profesional. Formatea el texto usando Markdown con títulos (##, ###), listas, y énfasis (**negrita**, *cursiva*).
"""

            client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'), timeout=30.0, max_retries=1)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto pedagogo que crea talleres educativos innovadores y efectivos. Responde siempre en formato Markdown."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            actividad_generada = response.choices[0].message.content
            
            # Convertir markdown a HTML
            import markdown
            actividad_html = markdown.markdown(actividad_generada, extensions=['nl2br', 'tables', 'fenced_code'])

            # Guardar en base de datos
            nueva_actividad = ActividadGenerada(
                user_id=current_user.id,
                nombre_profesor=nombre_profesor,
                grado=grado,
                asignatura=asignatura,
                tematica=tematica,
                cantidad_estudiantes=int(cantidad_estudiantes),
                tipo_actividad=tipo_actividad,
                tiempo=int(tiempo),
                recursos=recursos,
                contenido=actividad_generada
            )
            db.session.add(nueva_actividad)
            db.session.commit()
            
            # Guardar en sesión para poder exportar después
            session['ultima_actividad'] = {
                'id': nueva_actividad.id,
                'nombre_profesor': nombre_profesor,
                'grado': grado,
                'asignatura': asignatura,
                'tematica': tematica,
                'cantidad_estudiantes': cantidad_estudiantes,
                'tipo_actividad': tipo_actividad,
                'tiempo': tiempo,
                'recursos': recursos,
                'competencias': competencias,
                'contenido': actividad_generada
            }

            flash('Actividad generada y guardada exitosamente', 'success')
            return render_template('select_activities.html', 
                                 actividad=actividad_html,
                                 datos={
                                     'nombre_profesor': nombre_profesor,
                                     'grado': grado,
                                     'asignatura': asignatura,
                                     'tematica': tematica,
                                     'cantidad_estudiantes': cantidad_estudiantes,
                                     'tipo_actividad': tipo_actividad,
                                     'tiempo': tiempo,
                                     'recursos': recursos,
                                     'competencias': competencias
                                 })
        
        except Exception as e:
            flash(f'Error al generar actividad: {str(e)}', 'danger')
            return render_template('select_activities.html')

    # Modelos eliminados: Asignatura, Habilidad. Usar listas vacías.
    grados = Grado.query.all()
    asignaturas = []
    habilidades = []
    return render_template('select_activities.html', grados=grados, asignaturas=asignaturas, habilidades=habilidades)


@app.route('/mis_actividades')
@login_required
def mis_actividades():
    """Ver todas las actividades generadas por el usuario"""
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('Mis Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))

    actividades = ActividadGenerada.query.filter_by(user_id=current_user.id).order_by(ActividadGenerada.fecha_creacion.desc()).all()
    return render_template('mis_actividades.html', actividades=actividades)


@app.route('/ver_actividad/<int:actividad_id>')
@login_required
def ver_actividad(actividad_id):
    """Ver una actividad específica"""
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('Mis Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))
    actividad = ActividadGenerada.query.get_or_404(actividad_id)
    
    # Verificar que la actividad pertenece al usuario
    if actividad.user_id != current_user.id:
        flash('No tienes permiso para ver esta actividad', 'danger')
        return redirect(url_for('mis_actividades'))
    
    # Convertir markdown a HTML
    import markdown
    actividad_html = markdown.markdown(actividad.contenido, extensions=['nl2br', 'tables', 'fenced_code'])
    
    return render_template('ver_actividad.html', actividad=actividad, actividad_html=actividad_html)


@app.route('/editar_actividad/<int:actividad_id>', methods=['GET', 'POST'])
@login_required
def editar_actividad(actividad_id):
    """Editar una actividad existente"""
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('Mis Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))
    actividad = ActividadGenerada.query.get_or_404(actividad_id)
    
    # Verificar que la actividad pertenece al usuario
    if actividad.user_id != current_user.id:
        flash('No tienes permiso para editar esta actividad', 'danger')
        return redirect(url_for('mis_actividades'))
    
    if request.method == 'POST':
        try:
            actividad.nombre_profesor = request.form.get('nombre_profesor')
            actividad.grado = request.form.get('grado')
            actividad.asignatura = request.form.get('asignatura')
            actividad.tematica = request.form.get('tematica')
            actividad.cantidad_estudiantes = int(request.form.get('cantidad_estudiantes'))
            actividad.tipo_actividad = request.form.get('tipo_actividad')
            actividad.tiempo = int(request.form.get('tiempo'))
            actividad.recursos = request.form.get('recursos')
            actividad.contenido = request.form.get('contenido')
            actividad.fecha_modificacion = datetime.now()
            
            db.session.commit()
            flash('Actividad actualizada exitosamente', 'success')
            return redirect(url_for('ver_actividad', actividad_id=actividad.id))
        except Exception as e:
            db.session.rollback()
            flash(f'Error al actualizar actividad: {str(e)}', 'danger')
    
    return render_template('editar_actividad.html', actividad=actividad)


@app.route('/eliminar_actividad/<int:actividad_id>', methods=['POST'])
@login_required
def eliminar_actividad(actividad_id):
    """Eliminar una actividad"""
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('Mis Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))
    actividad = ActividadGenerada.query.get_or_404(actividad_id)
    
    # Verificar que la actividad pertenece al usuario
    if actividad.user_id != current_user.id:
        flash('No tienes permiso para eliminar esta actividad', 'danger')
        return redirect(url_for('mis_actividades'))
    
    try:
        db.session.delete(actividad)
        db.session.commit()
        flash('Actividad eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar actividad: {str(e)}', 'danger')
    
    return redirect(url_for('mis_actividades'))


@app.route('/exportar_actividad_word')
@app.route('/exportar_actividad_word/<int:actividad_id>')
@login_required
def exportar_actividad_word(actividad_id=None):
    """Exportar actividad a Word con formato profesional"""
    # Restringir acceso si el admin desactivó la función (excepto para admin)
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='generador_actividades_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('El Generador de Actividades está desactivado por el administrador.', 'warning')
        return redirect(url_for('dashboard'))
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from flask import send_file
    import re
    
    # Si se proporciona ID, buscar en BD
    if actividad_id:
        actividad_obj = ActividadGenerada.query.get_or_404(actividad_id)
        
        # Verificar permisos
        if actividad_obj.user_id != current_user.id:
            flash('No tienes permiso para exportar esta actividad', 'danger')
            return redirect(url_for('mis_actividades'))
        
        actividad = {
            'nombre_profesor': actividad_obj.nombre_profesor,
            'grado': actividad_obj.grado,
            'asignatura': actividad_obj.asignatura,
            'tematica': actividad_obj.tematica,
            'cantidad_estudiantes': str(actividad_obj.cantidad_estudiantes),
            'tipo_actividad': actividad_obj.tipo_actividad,
            'tiempo': str(actividad_obj.tiempo),
            'recursos': actividad_obj.recursos,
            'contenido': actividad_obj.contenido
        }
    else:
        # Usar sesión (última actividad generada)
        actividad = session.get('ultima_actividad')
        if not actividad:
            flash('No hay ninguna actividad para exportar', 'warning')
            return redirect(url_for('select_activities'))
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========== ENCABEZADO CON COLOR ==========
    # Título principal con fondo de color
    titulo = doc.add_heading(f"TALLER DE {actividad['asignatura'].upper()}", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.runs[0]
    titulo_run.font.size = Pt(24)
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
    titulo_run.font.bold = True
    
    # Agregar fondo azul al título
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '667eea')  # Color morado/azul
    titulo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # ========== TABLA DE INFORMACIÓN ==========
    # Construir datos de tabla dinámicamente
    
    # Configurar datos de la tabla
    datos_tabla = [
        ('👨‍🏫 Profesor', actividad['nombre_profesor']),
        ('📚 Grado', actividad['grado']),
        ('💡 Temática', actividad['tematica']),
        ('👥 Tipo', actividad['tipo_actividad']),
        ('🎓 Estudiantes', str(actividad['cantidad_estudiantes'])),
        ('⏱️ Tiempo', f"{actividad['tiempo']} minutos"),
        ('🛠️ Recursos', actividad['recursos'])
    ]
    # Agregar competencias si están presentes
    if 'competencias' in actividad and actividad['competencias']:
        comp_txt = actividad['competencias'] if isinstance(actividad['competencias'], str) else ", ".join(actividad['competencias'])
        datos_tabla.append(('🏆 Competencias', comp_txt))

    # Crear tabla con número de filas adecuado
    tabla_info = doc.add_table(rows=len(datos_tabla), cols=2)
    tabla_info.style = 'Light Grid Accent 1'
    
    for i, (campo, valor) in enumerate(datos_tabla):
        row_cells = tabla_info.rows[i].cells
        
        # Celda de campo (izquierda) con color
        cell_campo = row_cells[0]
        cell_campo.text = campo
        cell_campo.paragraphs[0].runs[0].font.bold = True
        cell_campo.paragraphs[0].runs[0].font.size = Pt(11)
        cell_campo.paragraphs[0].runs[0].font.color.rgb = RGBColor(68, 114, 196)
        
        # Agregar fondo gris claro
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E7E6E6')
        cell_campo.paragraphs[0]._element.get_or_add_pPr().append(shading)
        
        # Celda de valor (derecha)
        cell_valor = row_cells[1]
        cell_valor.text = valor
        cell_valor.paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ========== LÍNEA DECORATIVA ==========
    linea = doc.add_paragraph()
    linea_run = linea.add_run('═' * 80)
    linea_run.font.color.rgb = RGBColor(102, 126, 234)
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== CONTENIDO DE LA ACTIVIDAD ==========
    # Procesar el contenido markdown para Word
    contenido = actividad['contenido']
    
    # Procesar líneas del contenido
    lineas = contenido.split('\n')
    
    for linea in lineas:
        linea = linea.strip()
        if not linea:
            doc.add_paragraph()
            continue
        
        # Detectar títulos nivel 1 (#)
        if linea.startswith('# ') and not linea.startswith('## '):
            titulo_h1 = doc.add_heading(linea.replace('# ', '').strip(), level=1)
            titulo_h1.runs[0].font.color.rgb = RGBColor(68, 114, 196)
            titulo_h1.runs[0].font.size = Pt(18)
            
        # Detectar títulos nivel 2 (##)
        elif linea.startswith('## ') and not linea.startswith('### '):
            titulo_seccion = doc.add_heading(linea.replace('## ', '').strip(), level=2)
            titulo_seccion.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            titulo_seccion.runs[0].font.size = Pt(16)
            
        # Detectar títulos nivel 3 (###)
        elif linea.startswith('### ') and not linea.startswith('#### '):
            titulo_subseccion = doc.add_heading(linea.replace('### ', '').strip(), level=3)
            titulo_subseccion.runs[0].font.color.rgb = RGBColor(118, 75, 162)
            titulo_subseccion.runs[0].font.size = Pt(14)
            
        # Detectar títulos nivel 4 (####)
        elif linea.startswith('#### '):
            titulo_h4 = doc.add_heading(linea.replace('#### ', '').strip(), level=4)
            titulo_h4.runs[0].font.color.rgb = RGBColor(118, 75, 162)
            titulo_h4.runs[0].font.size = Pt(12)
            
        # Detectar listas con viñetas
        elif linea.startswith('- ') or linea.startswith('* '):
            texto_limpio = linea[2:].strip()
            # Procesar negritas en listas
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            
            # Procesar formato inline
            partes = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', texto_limpio)
            for parte in partes:
                if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
                    run = p.add_run(parte[2:-2])
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(233, 30, 99)
                    run.font.size = Pt(11)
                elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2:
                    run = p.add_run(parte[1:-1])
                    run.font.italic = True
                    run.font.size = Pt(11)
                elif parte:
                    run = p.add_run(parte)
                    run.font.size = Pt(11)
            
        # Detectar listas numeradas
        elif re.match(r'^\d+\.\s', linea):
            texto_lista = re.sub(r'^\d+\.\s', '', linea).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            
            # Procesar formato inline
            partes = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', texto_lista)
            for parte in partes:
                if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
                    run = p.add_run(parte[2:-2])
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(233, 30, 99)
                    run.font.size = Pt(11)
                elif parte.startswith('*') and parte.endswith('*') and len(parte) > 2:
                    run = p.add_run(parte[1:-1])
                    run.font.italic = True
                    run.font.size = Pt(11)
                elif parte:
                    run = p.add_run(parte)
                    run.font.size = Pt(11)
            
        # Párrafo normal
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            
            # Procesar formato inline (negritas, cursivas)
            # Primero procesar negritas con cursiva (***texto***)
            partes = re.split(r'(\*\*\*[^*]+\*\*\*)', linea)
            
            for parte in partes:
                if parte.startswith('***') and parte.endswith('***'):
                    # Negrita + Cursiva
                    run = p.add_run(parte[3:-3])
                    run.font.bold = True
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(233, 30, 99)
                    run.font.size = Pt(11)
                else:
                    # Seguir procesando negritas y cursivas simples
                    sub_partes = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', parte)
                    
                    for sub_parte in sub_partes:
                        if sub_parte.startswith('**') and sub_parte.endswith('**') and len(sub_parte) > 4:
                            # Negrita
                            run = p.add_run(sub_parte[2:-2])
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(233, 30, 99)
                            run.font.size = Pt(11)
                        elif sub_parte.startswith('*') and sub_parte.endswith('*') and len(sub_parte) > 2:
                            # Cursiva
                            run = p.add_run(sub_parte[1:-1])
                            run.font.italic = True
                            run.font.size = Pt(11)
                        elif sub_parte:
                            # Texto normal
                            run = p.add_run(sub_parte)
                            run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ========== PIE DE PÁGINA ==========
    linea_footer = doc.add_paragraph()
    linea_footer_run = linea_footer.add_run('═' * 80)
    linea_footer_run.font.color.rgb = RGBColor(102, 126, 234)
    linea_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer = doc.add_paragraph()
    footer_run1 = footer.add_run('✨ Generado por Plataforma ADAT ✨')
    footer_run1.font.bold = True
    footer_run1.font.color.rgb = RGBColor(102, 126, 234)
    footer_run1.font.size = Pt(10)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fecha_footer = doc.add_paragraph()
    fecha_run = fecha_footer.add_run(f"📅 {datetime.now().strftime('%d de %B de %Y - %I:%M %p')}")
    fecha_run.font.italic = True
    fecha_run.font.size = Pt(9)
    fecha_run.font.color.rgb = RGBColor(128, 128, 128)
    fecha_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Guardar en memoria
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Taller_{actividad['asignatura']}_{fecha}.docx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


# ==================== MIS RESULTADOS (RESUMEN) ====================
@app.route('/mis_resultados')
@login_required
def mis_resultados():
    """Vista resumen con el estado de resultados del usuario en todos los tests"""
    # Flag de acceso a Resultados
    try:
        conf = ConfiguracionSistema.query.filter_by(clave='resultados_activo').first()
        activo = (conf.valor == 'true') if conf else True
    except Exception:
        activo = True
    if current_user.rol != 'admin' and not activo:
        flash('La vista de Resultados está desactivada por el administrador.', 'warning')
        return redirect(url_for('dashboard'))
    r1 = ResultadoQuiz.query.filter_by(user_id=current_user.id).first()
    r2 = ResultadoQuizDos.query.filter_by(user_id=current_user.id).first()
    r3 = ResultadoQuizTres.query.filter_by(user_id=current_user.id).first()
    r4 = ResultadoQuizCuatro.query.filter_by(user_id=current_user.id).first()
    return render_template('mis_resultados.html', r1=r1, r2=r2, r3=r3, r4=r4)
